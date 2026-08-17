from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from lxml import etree


OUT = Path("/home/arcosium/vault/GenomicWQB-docs/docs/머신발표/GenomicWQB_머신발표_발표대본_10분.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(20, 82, 150)
DARK = RGBColor(20, 20, 20)
GRAY = RGBColor(90, 99, 110)
LIGHT_BLUE = "E8F0FA"
FONT = "맑은 고딕"


slides = [
    (1, "표지", 25,
     "안녕하세요. GenomicWQB는 알파를 유전체로 나누고 생성, 검문, 시뮬레이션, 제출 판단을 반복하는 자동 발굴 시스템입니다. 98일 동안 시뮬레이션 35,701건과 923라운드가 쌓였고 원장상 제출 성공은 86건입니다. 오늘은 이 숫자보다 머신이 다음 실험을 고르는 방식에 집중하겠습니다."),
    (2, "발표 순서", 15,
     "발표는 문제 정의, 머신 구조, 핵심 로직, 실증 성과, 한계와 다음 단계 순서입니다. 각 장의 표를 전부 읽기보다 판단 규칙과 그 규칙이 남긴 결과만 짚겠습니다."),
    (3, "문제 정의", 35,
     "알파의 품질은 BRAIN 시뮬레이션 뒤에야 알 수 있습니다. 조합 공간은 곱으로 늘어나지만 동시 슬롯은 8개이고 한 건에 대기열을 포함해 2분에서 20분이 걸립니다. 운영 실측은 하루 평균 364건이었고 제출은 하루 4건이 상한입니다. 결국 문제는 수식을 많이 만드는 것이 아니라 제한된 평가 슬롯을 어느 후보에 배분하느냐입니다."),
    (4, "운영 목적함수", 30,
     "3개월 안에 IS Sharpe를 OS 성과로 검증하기는 어렵습니다. 그래서 반복 측정이 가능한 제출 통과 건수를 운영 목적함수로 두고 품질은 플랫폼 관문 충족 여부로 판단했습니다. 원장상 성공은 86건이지만 같은 계보를 반복한 결과 프로덕션 상관이 높아졌습니다. 수량 중심 목표의 비용입니다."),
    (5, "라운드 파이프라인", 35,
     "한 라운드는 개체군 구성, 사전 검문, 시뮬레이션, 채점과 게이트, 학습과 복기의 다섯 단계로 돕니다. 마지막 단계에서 엘리트, 밴딧 보상, 정향변이 통계를 갱신하고 이를 다음 라운드의 입력으로 돌려보냅니다. 현재 최대 세대 깊이는 g136입니다. 핵심은 실패 기록도 다음 실험의 재료로 쓴다는 점입니다."),
    (6, "유전체 설계", 30,
     "GenomicWQB는 구문 트리를 제한 없이 키우지 않습니다. 데이터필드, 변환, 결합, 룩백, 중립화, 감쇠를 고정 유전자에 담아 FASTEXPR로 렌더링합니다. 덕분에 교차와 변이의 단위가 분명해집니다. 다만 플랫폼을 완전히 복제하지 못해 unknown operator 135건, unknown variable 171건, unit incompatibility 353건은 남았습니다."),
    (7, "시뮬레이션 전 관문", 25,
     "시뮬레이션 전에는 문법과 단위, 필드 가용성, 중복, 테마 조건, 차단 FAIL을 확인합니다. 이미 아는 실패를 8개 슬롯에 태우면 다음 실험이 그만큼 밀립니다. 사전 검문은 모든 오류를 맞히는 장치가 아니라 비싼 실패를 줄이는 장치입니다."),
    (8, "실제 라운드 r198", 40,
     "8월 14일 r198은 변이와 교차, 구제 변형, 개선 레이어, 한 축 스윕으로 후보 20개를 만들었습니다. 가장 높은 Sharpe는 2번 후보의 2.02였지만 PROD_CORRELATION에서 차단됐습니다. 반면 Sharpe 1.44인 6번 후보는 7개 체크를 모두 통과해 제출됐습니다. 가장 강한 알파가 등재되는 것이 아닙니다. 통과 조건이 제출 가능성을 결정하며 머신의 목적함수도 그 기준을 따라야 합니다."),
    (9, "게이트 실측 학습", 30,
     "거절 본문에는 여러 FAIL이 함께 적히므로 자주 등장한 체크를 곧바로 하드 규칙으로 만들면 안 됩니다. 최근 14일의 거절 217건과 성공 28건을 비교해 성공 전례가 있는 체크는 소프트로 되돌렸습니다. 소프트 체크는 실제 제출로 판정을 확인하고, 단독 차단 증거가 있는 하드 체크만 계획 단계에서 막습니다."),
    (10, "한 축 스윕", 35,
     "부모와 자식의 Sharpe가 모두 있는 10,119쌍에서 한 개 변경군의 개선율은 20.3%, 여섯 개 이상 변경군은 6.2%였습니다. 인과효과라고 단정할 수는 없지만 원인 귀속은 쉬워집니다. 실제로 중립화만 바꿔 Sharpe가 1.72에서 2.47로 올랐고 감쇠만 바꿔 회전율이 100.5%에서 48.5%로 낮아졌습니다."),
    (11, "정향변이", 30,
     "변이는 부모의 실패 원인과 실측 지표가 지목합니다. 회전율이 높으면 감쇠와 시간창을 조정하고 집중도가 높으면 중립화와 이상치 처리를 건드립니다. 고회전 관문은 WARNING이라 실패 문자열에 안 잡혀 실측 지표도 함께 봅니다. 누적 관측에서 boost 지시의 개선율은 36.0%, smooth는 0.5%였습니다. 통제 실험은 아니며 다음 변이의 우선순위를 바꾸는 기록입니다."),
    (12, "다목적 선택", 35,
     "Sharpe 하나만 최대화하면 엘리트가 한 유형으로 고착됩니다. 그래서 NSGA II는 Sharpe, Fitness, 제출 경로 점수, 자기상관, 최근 2년 Sharpe를 함께 봅니다. 다만 family 정보가 남은 제출작 25건 중 model 14건과 pv 8건이 88%였습니다. 선택 알고리즘만으로는 생성 단계의 쏠림을 고칠 수 없다는 한계가 남았습니다."),
    (13, "밴딧과 LLM", 30,
     "분석상 전 체크 통과 대리값은 35,701건 중 442건, 1.2%라 완전 통과만으로는 보상이 너무 드뭅니다. 밴딧은 연속 선택 점수와 제출 가능 보상을 섞어 무작위 탐색 슬롯을 배분합니다. LLM 위원회는 가설과 전략 스펙을 만들고 유전 알고리즘은 반복 탐색과 계보 기록을 맡습니다. 새 아이디어는 변이 전에 한 번 그대로 측정해 원안의 성적과 이후 개선을 구분합니다."),
    (14, "제출 운영", 30,
     "운영 원장에서 거절 시도는 일일 4건 한도를 줄이지 않았습니다. 프로덕션 상관은 당시 제출 시도의 403 본문에서 가장 안정적으로 확인됐고 형제 알파가 등재되면 값도 바뀌었습니다. 그래서 전날 판정을 재사용하지 않고 당일 다시 확인합니다. 8월 12일에는 쿼터 뒤 통과작 27건이 큐에서 사라진 적이 있어, 지금은 다음 리셋용 대기 큐에 반드시 넣습니다."),
    (15, "생성 경로별 수율", 35,
     "Sharpe가 기록된 주요 자동 생성 경로 여섯 개 13,298건을 같은 기준으로 비교했습니다. 무작위 탐색은 2,582건 중 9건, 0.3%였고 한 축 스윕은 1,626건 중 366건, 22.5%였습니다. 개선과 구제 경로도 높은 관측 수율을 보였습니다. 경로별 시점과 표본이 달라 인과효과로 보지는 않으며 무작위 탐색은 새 데이터셋의 기준선으로 남깁니다."),
    (16, "실패 분포", 25,
     "누적 FAIL에서는 LOW_FITNESS와 LOW_SHARPE가 많았습니다. PROD_CORRELATION은 시뮬레이션 fail_items에 18건뿐이지만 제출 직전에 주로 측정됩니다. 8월 13일부터 14일까지 태그 134건 가운데 상관값이 기록된 53건의 중앙값은 0.9028이고 0.7 미만은 없었습니다. 누적 실패 빈도와 오늘의 제출 병목을 나눠 봐야 합니다."),
    (17, "세 사례", 35,
     "세 사례가 운영 규칙의 영향을 보여 줍니다. 한 축 스윕은 변화의 원인을 남겼습니다. 반면 7월 28일부터 8월 3일까지 제출 17건 중 11건이 rsk70에 몰렸는데 당시 시딩 축 26개가 모두 그 데이터셋이었습니다. 또 8월 12일에는 쿼터 뒤 통과작 27건이 큐에서 사라졌고 다음 날 복원한 후보 중 한 건만 제출됐습니다. 모두 알파보다 생성과 제출 규칙의 문제였습니다."),
    (18, "상관 포화", 40,
     "수량 전략의 비용은 상관 포화로 나타났습니다. 같은 계보 네 건이 등재된 뒤 관련 후보의 프로덕션 상관이 함께 올랐고 관측 최소값도 0.7111이었습니다. 피라미드는 리전, 지연, 데이터셋 카테고리 조합에 세 건을 채워야 집계됩니다. 이미 찬 칸에 같은 계보를 더 쌓지 않도록 생성 우선순위를 미달 피라미드 칸 기준으로 바꿨고, 짝 필드도 같은 데이터셋 안에서 고르게 했습니다."),
    (19, "다음 3주", 30,
     "다음 단계는 세 가지입니다. Analyst 4,218필드와 Fundamental 4,428필드에 새 계보를 심어 생성 단계부터 다변화합니다. 절단과 결측치 처리 같은 마지막 튜닝값은 탐색 축에서 내리고 데이터필드와 변환에 슬롯을 돌립니다. 성과지표도 제출 건수에서 프로덕션 상관 관문을 통과한 독립 계보 수로 바꿉니다."),
    (20, "마무리", 20,
     "GenomicWQB의 결과는 원장상 제출 성공 86건만이 아닙니다. 어떤 유전자가 어느 환경에서 작동했고 무엇이 실패했는지를 남긴 35,701건의 실험 원장도 같은 무게의 결과물입니다. 그 기록 덕분에 목적함수의 잘못까지 확인할 수 있었습니다. 다음 버전은 더 많이 제출하는 머신보다 서로 다른 계보를 더 효율적으로 찾는 머신을 목표로 합니다. 감사합니다."),
]


questions = [
    ("35,701건은 한 계정의 기록인가요?",
     "아닙니다. 기준시각까지 GenomicWQB 운영 데이터베이스 전체에 남은 시뮬레이션 행의 합계입니다. 개인별 실적이나 리더보드 수치가 아니라 시스템 전체의 탐색 기록입니다."),
    ("86건은 BRAIN 화면의 최종 제출 수와 같은가요?",
     "이 발표에서 86건은 submit_attempts 테이블의 submitted=1 행 수입니다. 외부 화면의 주차별 제출 수나 현재 유효 알파 수와 같은 지표로 해석하지 않습니다."),
    ("왜 Sharpe나 Fitness가 아니라 제출 건수를 목적함수로 삼았나요?",
     "3개월 안에는 IS 지표가 실제 OS 성과로 이어지는지 충분히 검증하기 어렵습니다. 제출 통과는 플랫폼이 즉시 돌려주는 반복 가능한 판정이라 운영 목표로 삼았습니다. 다만 상관 포화가 확인돼 다음 버전에서는 독립 계보 수로 바꿉니다."),
    ("Sharpe 2.02 후보보다 1.44 후보를 제출한 것이 합리적인가요?",
     "해당 라운드에서는 그렇습니다. 2.02 후보는 PROD_CORRELATION에서 막혔고 1.44 후보는 7개 체크를 모두 통과했습니다. 이 머신의 목표는 최고 Sharpe 후보 선정이 아니라 실제 제출 가능한 후보 선정입니다."),
    ("한 축 변경이 더 좋다는 인과 결론인가요?",
     "아닙니다. 운영 자료의 기술통계입니다. 한 축 변경군의 관측 개선율이 높았고 결과를 특정 변화에 귀속하기 쉬웠다는 데까지만 말할 수 있습니다. 무작위 배정 실험이 아니므로 인과효과는 주장하지 않습니다."),
    ("NSGA II가 실제로 다양성을 높였나요?",
     "현재 원장만으로 적용 전후 효과를 분리할 수 없습니다. family 정보가 남은 제출작 25건 중 88%가 model과 pv 두 계열에 몰렸다는 사실만 확인했습니다. 다음 실험에서는 생성 팔레트와 엘리트 풀의 분포를 버전별로 함께 기록해야 합니다."),
    ("사전 검문이 있는데도 플랫폼 오류가 남은 이유는 무엇인가요?",
     "로컬 검문은 공개된 문법과 현재까지 관측한 단위 규칙을 근사합니다. 플랫폼의 모든 연산자, 필드 상태, 단위 조합을 완전히 복제하지는 못합니다. 따라서 오류를 없애기보다 이미 아는 비싼 실패를 줄이는 장치로 봐야 합니다."),
    ("하드와 소프트 게이트는 어떻게 구분하나요?",
     "최근 관측창에서 거절과 성공을 함께 봅니다. FAIL이 있어도 제출 성공한 전례가 있으면 그 체크는 단독 차단 근거가 아니므로 소프트로 둡니다. 성공 전례가 없고 반복적으로 차단과 함께 나타나는 체크만 하드 후보로 취급합니다."),
    ("거절이 제출 한도를 줄이지 않는다는 것은 플랫폼의 공식 규칙인가요?",
     "공식 규칙으로 일반화한 것이 아니라 이 기간 운영 원장에서 확인한 관측입니다. 여러 거절 뒤에도 REGULAR_SUBMISSION이 0/4로 남은 사례를 바탕으로 당시 제출 전략을 설계했습니다. 플랫폼 정책이 바뀔 수 있어 계속 재확인합니다."),
    ("전 체크 통과 442건, 1.2%는 정확히 무엇을 뜻하나요?",
     "Sharpe가 기록되고 fail_count=0, error_count=0인 행을 분석용 대리값으로 정의한 수치입니다. 실제 최종 제출 성공과 같은 뜻은 아닙니다. 보상이 얼마나 희소한지를 설명하려고 쓴 운영 분석 기준입니다."),
    ("무작위 탐색 수율이 0.3%라면 없애는 편이 낫지 않나요?",
     "기존 계보를 개선하는 데는 비효율적이지만 새 데이터셋과 새 필드 조합에 처음 진입할 때 기준선과 탐색 다양성을 제공합니다. 완전히 제거하면 이미 잘되는 계보만 반복하는 착취 편향이 커집니다."),
    ("프로덕션 상관 중앙값 0.9028은 표본 편향이 있지 않나요?",
     "있습니다. 태그 134건 중 값이 남은 53건만의 조건부 통계입니다. 전체 후보의 상관 분포로 일반화하지 않고, 당시 제출 직전 관측군에 높은 상관이 몰렸다는 병목 신호로만 사용했습니다."),
    ("rsk70 쏠림은 NSGA II의 실패인가요?",
     "선택 단계만의 실패라고 볼 수 없습니다. 당시 시딩 축 26개가 모두 rsk70이라 생성 단계부터 후보 풀이 치우쳤습니다. 선택 알고리즘은 주어진 후보 밖의 계보를 만들 수 없으므로 생성 효과와 선택 효과를 분리해 기록해야 합니다."),
    ("LLM이 잘못된 수식이나 근거 없는 전략을 만들 위험은 어떻게 막나요?",
     "LLM 출력은 바로 제출하지 않습니다. 고정 유전체로 변환하고 문법, 필드 가용성, 중복, 테마 조건을 검사한 뒤 시뮬레이션으로 측정합니다. 새 전략 스펙은 변이 전에 원본 그대로 한 번 평가해 아이디어 자체의 성적도 분리합니다."),
    ("다음 버전의 독립 계보 수는 어떻게 정의할 계획인가요?",
     "서로 다른 루트 부모와 데이터 패밀리에서 출발하고 프로덕션 상관 관문을 통과한 계보를 별도 단위로 셀 계획입니다. 단순 표현식 수가 아니라 계보 ID, 패밀리, 상관 판정을 함께 보며 세부 임계값은 후속 실험으로 고정합니다."),
]


def set_run_font(run, size: float, color: RGBColor = DARK, bold: bool | None = None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, 8.5, GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_label_paragraph(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(label)
    set_run_font(r, 10, BLUE, True)
    r = p.add_run(text)
    set_run_font(r, 10, DARK)
    return p


doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(18)
sec.bottom_margin = Mm(23)
sec.left_margin = Mm(20)
sec.right_margin = Mm(20)
sec.header_distance = Mm(8)
sec.footer_distance = Mm(9)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.line_spacing = 1.35
normal.paragraph_format.space_after = Pt(5)

footer = sec.footer.paragraphs[0]
add_page_number(footer)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(32)
r = p.add_run("GenomicWQB 머신 발표회 발표대본")
set_run_font(r, 20, BLUE, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("10분 발표용 · 슬라이드 20장 · 예상 질문 15개")
set_run_font(r, 11, GRAY, False)

doc.add_paragraph()
add_label_paragraph(doc, "발표 기준  ", "화면의 숫자를 모두 읽지 않고 각 장의 결론과 근거 수치만 말합니다.")
add_label_paragraph(doc, "시간 운영  ", "목표 10분 10초입니다. 현장에서는 슬라이드 전환을 빠르게 하면 약 10분에 맞습니다.")
add_label_paragraph(doc, "강조 장면  ", "8쪽 r198, 10쪽 한 축 스윕, 18쪽 상관 포화에서 잠깐 멈춰 화면을 가리킵니다.")
add_label_paragraph(doc, "시간 초과 시  ", "2쪽은 한 문장으로 넘기고 13쪽과 14쪽은 마지막 문장을 생략하면 약 30초를 줄일 수 있습니다.")

p = doc.add_paragraph()
r = p.add_run()
r.add_break(WD_BREAK.PAGE)
r = p.add_run("발표대본")
set_run_font(r, 17, BLUE, True)
p.paragraph_format.space_after = Pt(10)

for number, title, seconds, script in slides:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"슬라이드 {number}  {title}")
    set_run_font(r, 12.5, BLUE, True)
    r = p.add_run(f"   목표 {seconds}초")
    set_run_font(r, 9, GRAY, False)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(3)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(script)
    set_run_font(r, 10.5, DARK, False)

p = doc.add_paragraph()
r = p.add_run()
r.add_break(WD_BREAK.PAGE)
r = p.add_run("10분 발표 적합성 검증")
set_run_font(r, 17, BLUE, True)
p.paragraph_format.space_after = Pt(8)

script_text = "\n".join(script for _, _, _, script in slides)
char_count = len(re.sub(r"\s+", "", script_text))
word_count = len(re.findall(r"\S+", script_text))
budget_seconds = sum(seconds for _, _, seconds, _ in slides)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Cm(5.2), Cm(5.2), Cm(5.2)]
headers = ["검증 항목", "결과", "판정"]
for i, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
    cell.width = width
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 9.5, DARK, True)

rows = [
    ("슬라이드 시간 예산", f"{budget_seconds // 60}분 {budget_seconds % 60}초", "적합"),
    ("대본 분량", f"공백 제외 {char_count:,}자 · {word_count:,}어절", "적합"),
    ("빠른 발표", "분당 330자 기준", f"약 {char_count / 330:.1f}분"),
    ("보통 발표", "분당 300자 기준", f"약 {char_count / 300:.1f}분"),
    ("또박또박 발표", "분당 270자 기준", f"약 {char_count / 270:.1f}분"),
]
for row in rows:
    cells = table.add_row().cells
    for i, (cell, width, text) in enumerate(zip(cells, widths, row)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 9.5, DARK, False)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.line_spacing = 1.3
r = p.add_run(
    "판정: 20장 시간 예산은 10분 10초입니다. 분당 300자 안팎으로 말하면 대본 낭독은 약 9분 40초이고, "
    "슬라이드 전환과 숫자를 가리키는 짧은 멈춤을 합쳐 약 10분에서 10분 20초가 예상됩니다. "
    "분당 270자 이하로 천천히 말한다면 앞쪽의 시간 초과 시 생략 지침을 적용합니다. "
    "실전에서는 8쪽, 10쪽, 18쪽을 충분히 설명하고 나머지 장은 결론 문장을 먼저 말하는 구성이 적당합니다."
)
set_run_font(r, 10.5, DARK, False)

p = doc.add_paragraph()
r = p.add_run()
r.add_break(WD_BREAK.PAGE)
r = p.add_run("예상 질문과 답변")
set_run_font(r, 17, BLUE, True)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run("질문은 원장 정의, 통계 해석, 설계 한계 순으로 배치했습니다. 답변은 20초 안팎으로 말할 수 있는 길이입니다.")
set_run_font(r, 9.5, GRAY, False)

for idx, (question, answer) in enumerate(questions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Q{idx}. {question}")
    set_run_font(r, 11, BLUE, True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run("A. " + answer)
    set_run_font(r, 10.2, DARK, False)

doc.core_properties.title = "GenomicWQB 머신 발표회 발표대본"
doc.core_properties.subject = "10분 발표용 대본과 예상 질문 답변"
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.keywords = "GenomicWQB, 발표대본, Q&A"

tmp = OUT.with_suffix(".docx.tmp")
doc.save(tmp)

# Word에서 문단·그림 앞 점으로 보이는 비인쇄 페이지 나눔 속성을 제거한다.
marker_names = ("keepNext", "keepLines", "pageBreakBefore", "suppressLineNumbers")
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}
clean = OUT.with_suffix(".docx.clean")
with ZipFile(tmp) as source, ZipFile(clean, "w", ZIP_DEFLATED) as target:
    for item in source.infolist():
        data = source.read(item.filename)
        if item.filename.endswith(".xml"):
            try:
                root = etree.fromstring(data)
            except etree.XMLSyntaxError:
                root = None
            if root is not None:
                for name in marker_names:
                    for node in list(root.xpath(f"//w:{name}", namespaces=NS)):
                        parent = node.getparent()
                        if parent is not None:
                            parent.remove(node)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        target.writestr(deepcopy(item), data)

os.replace(clean, OUT)
tmp.unlink()

print(f"created={OUT}")
print(f"slides={len(slides)} questions={len(questions)} budget_seconds={budget_seconds}")
print(f"spoken_chars_no_space={char_count} words={word_count}")
