from pathlib import Path

from docx import Document


DOCX = Path('/home/arcosium/projects/GenomicWQB/docs/유전알고리즘_알파리서치_리포트.docx')


def main():
    doc = Document(DOCX)
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace(
        '다양성은 선택뿐 아니라 생성 단계에서도 확보해야 한다.',
        '다양성은 선택뿐 아니라 생성 단계에서도 확보할 필요가 있다.',
    )
    doc.paragraphs[22].text = (
        '탐색 설계에서 물을 수 있는 것은 많지만 검증 가치가 있는 질문에는 두 조건이 있다. '
        '답은 자명하지 않고 그 주의 플랫폼 조건이 바뀌어도 결론이 뒤집히지 않아야 한다. '
        '이 기준으로 두 질문을 남겼다. 하나는 변이의 폭, 다른 하나는 선택의 목적축을 다룬다.'
    )
    doc.paragraphs[60].text = doc.paragraphs[60].text.replace(
        'RQ1에 대해서는 두 가지를 구분해 답해야 한다.',
        'RQ1은 두 가지로 나눠 답한다.',
    )
    doc.paragraphs[87].text = doc.paragraphs[87].text.replace(
        '다음 실험에서는 생성 팔레트뿐 아니라 엘리트 풀의 분포와 교체율도 함께 기록해야 한다.',
        '다음 실험에서는 생성 팔레트뿐 아니라 엘리트 풀의 분포와 교체율도 함께 기록한다.',
    )
    doc.save(DOCX)


if __name__ == '__main__':
    main()
