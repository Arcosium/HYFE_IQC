from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


ROOT = Path('/home/arcosium/projects/GenomicWQB')
DOCX = ROOT / 'docs/유전알고리즘_알파리서치_리포트.docx'
FIGS = ROOT / 'docs/머신발표/figs'


def main():
    doc = Document(DOCX)
    figures = {
        43: ('fig1_genes_vs_gain.png', Mm(105)),
        48: ('fig2_origin_yield.png', Mm(144)),
        53: ('fig4_directives.png', Mm(140)),
        66: ('fig6_family_skew.png', Mm(94)),
    }
    for idx, (name, width) in figures.items():
        p = doc.paragraphs[idx]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = 0
        p.add_run().add_picture(str(FIGS / name), width=width)
    doc.save(DOCX)


if __name__ == '__main__':
    main()
