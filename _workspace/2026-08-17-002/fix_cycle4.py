from pathlib import Path

from docx import Document


ROOT = Path('/home/arcosium/projects/GenomicWQB')
DOCX = ROOT / 'docs/유전알고리즘_알파리서치_리포트.docx'


def main():
    doc = Document(DOCX)

    # 그림 캡션은 앞 그림과만 묶고, 뒤 본문까지 한 덩어리로 밀지 않는다.
    for idx in (44, 49, 54, 67):
        doc.paragraphs[idx].paragraph_format.keep_with_next = False
        doc.paragraphs[idx].paragraph_format.keep_together = True

    # 표 제목은 표 앞에 둬 다음 쪽 첫 줄에 홀로 남지 않게 한다.
    for p_idx, table_idx in ((50, 2), (58, 3), (76, 4)):
        caption = doc.paragraphs[p_idx]
        caption.paragraph_format.keep_with_next = True
        doc.tables[table_idx]._tbl.addprevious(caption._p)

    for p in doc.paragraphs:
        fixed = p.text.replace('`[13시전 수동]` 11건과 `[13시전 블라스트]` 123건',
                               '[13시전 수동] 11건과 [13시전 블라스트] 123건')
        if fixed != p.text:
            p.text = fixed

    doc.save(DOCX)


if __name__ == '__main__':
    main()
