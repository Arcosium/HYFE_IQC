from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


ROOT = Path('/home/arcosium/projects/GenomicWQB')
DOCX = ROOT / 'docs/유전알고리즘_알파리서치_리포트.docx'
FIGS = ROOT / 'docs/머신발표/figs'


def main():
    doc = Document(DOCX)

    # python-docx 저장 과정에서 유실된 기존 인라인 차트 네 개를 원본 자산으로 복원한다.
    figures = {
        43: ('fig1_genes_vs_gain.png', Mm(144)),
        48: ('fig2_origin_yield.png', Mm(144)),
        53: ('fig4_directives.png', Mm(140)),
        66: ('fig6_family_skew.png', Mm(94)),
    }
    for idx, (name, width) in figures.items():
        p = doc.paragraphs[idx]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = 0
        p.add_run().add_picture(str(FIGS / name), width=width)

    for p in doc.paragraphs:
        fixed = p.text.replace('`[13시전 수동]`·`[13시전 블라스트]`', '[13시전 수동]·[13시전 블라스트]')
        fixed = fixed.replace('따른 것이다[9].', '따른 것이다[8].')
        if fixed != p.text:
            p.text = fixed

    doc.save(DOCX)


if __name__ == '__main__':
    main()
