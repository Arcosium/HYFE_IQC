from pathlib import Path

from docx import Document


DOCX = Path(
    "/home/arcosium/projects/GenomicWQB/docs/유전알고리즘_알파리서치_리포트.docx"
)


def set_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def set_cell(cell, text: str) -> None:
    set_paragraph(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph(paragraph, "")


def main() -> None:
    document = Document(DOCX)
    replacements = {
        6: (
            "알파 리서치는 수식을 잘 쓰는 문제가 아니라 평가 예산을 어디에 던질지 정하는 문제다. "
            "후보의 품질은 WorldQuant BRAIN 시뮬레이션을 돌린 뒤에야 드러나고 동시 실행 슬롯은 "
            "여덟 개라 전수 탐색은 성립하지 않는다. 본 연구는 알파를 고정 유전자 공간의 유전체로 "
            "표현하고 선택·교차·변이를 반복하는 자동 발굴 시스템 GenomicWQB를 98일간 운영한 원장을 "
            "분석한다. 시뮬레이션 35,701건, 923라운드, 고유 표현식 26,404종이 분석 대상이다."
        ),
        10: (
            "본 시스템은 알파의 사후 수익률이 아니라 제출 통과 건수를 운영 목표로 삼았다. 2026년 "
            "8월 14일 16시 20분까지 submit_attempts에는 submitted=1인 기록 86건이 남았다. 그 선택이 "
            "남긴 비용도 함께 본다. 8월 13~14일 원장의 [13시전 수동]·[13시전 블라스트] 태그 134건 "
            "가운데 프로덕션 상관값이 기록된 53건의 중앙값은 0.9028, 최소값은 0.7111이었다. 관측값 "
            "중 0.7 아래는 한 건도 없었다."
        ),
        16: (
            "문제는 그 조합의 품질을 미리 알 수 없다는 데 있다. 후보 하나의 성적은 BRAIN "
            "시뮬레이션을 돌린 뒤에야 나오고 한 건에 대기열을 포함해 2분에서 20분이 걸린다. 계정당 "
            "동시 실행 슬롯은 여덟 개다. 대기열과 호출 제한, 실패 재시도를 포함한 98일 운영 실측은 "
            "하루 평균 364건이었다. 조합 공간과 평가 예산의 격차가 커서 전수 탐색은 논의 대상이 아니다."
        ),
        17: (
            "그래서 이 연구는 알파 리서치를 수식 설계 문제가 아니라 예산 배분 문제로 본다. 질문은 "
            "\"어떤 수식이 좋은가\"가 아니라 \"다음 여덟 슬롯을 어디에 던질 것인가\"이다. 유전 "
            "알고리즘은 넓은 탐색 공간을 재현 가능한 실험 단위로 나누고 성과가 좋은 구조를 다음 "
            "세대로 넘기며 실패에서도 다음 실험 방향을 얻는 절차를 제공한다[2,3]."
        ),
        19: (
            "이 시스템은 목적함수를 알파의 사후 수익률이 아니라 제출 통과 건수로 잡았다. 3개월 "
            "부트캠프 안에서는 IS Sharpe를 OS 성과로 검증하기 어렵지만 제출 통과는 플랫폼이 그날 "
            "돌려주는 판정이라 즉시 반복 측정할 수 있다. 같은 기간 원장에는 submit_attempts의 "
            "submitted=1 기록 86건이 남았고 상관 포화도 함께 나타났다."
        ),
        26: (
            "유전 프로그래밍과는 구분된다[2]. 유전 프로그래밍은 구문 트리 구조 자체를 진화시키지만 "
            "GenomicWQB는 문법과 연구 관행을 반영한 고정 유전자 공간에서 수식을 렌더링한다. 표현 "
            "자유도를 일부 포기해 오류와 복잡도를 줄였지만 플랫폼 규칙을 완전히 복제하지는 못했다. "
            "35,701건 가운데 WQB 응답에는 unknown operator 135건, unknown variable 171건, unit "
            "incompatibility 353건이 남았다. 사전 검문은 오류를 없애는 장치가 아니라 비싼 실패를 줄이는 "
            "장치다."
        ),
        28: (
            "이 문제의 특수성은 적합도 평가가 비싸고 외부에 있다는 데 있다. 평가자는 우리가 통제하지 "
            "않는 블랙박스이며 주간 테마와 일부 체크 조건도 시점에 따라 달라진다. 평가 결과가 다른 "
            "개체의 평가에 영향을 주기도 한다. 형제 알파가 등재되면 같은 계보의 프로덕션 상관이 수 "
            "시간 안에 올라 그 계보의 제출 가능성이 떨어질 수 있다. 적합도 함수가 고정되어 있지 않고 "
            "개체군의 역사에 의존한다는 뜻이다. 표준적인 유전 알고리즘 문헌이 다루는 정적 적합도와 "
            "다른 지점이며 뒤에서 살펴볼 상관 포화 현상과도 이어진다."
        ),
        31: (
            "알파 하나는 모델·패밀리·세대 메타데이터를 제외한 스물두 개 탐색 유전자로 표현된다. "
            "성격에 따라 세 갈래로 나뉜다."
        ),
        32: (
            "BRAIN 문서는 37일이나 14일 같은 임의값 대신 5·20·60·120·252일처럼 단순하고 합리적인 "
            "값을 권한다[8]. GenomicWQB의 주요 변이 경로는 여기에 2·3·10일을 더한 여덟 개 창을 "
            "우선한다. 다만 일부 무작위·세 번째 팩터 경로에는 40일이 남아 있고 원장에도 과거의 "
            "비표준 창이 존재한다. 따라서 표준 창 제약은 주된 탐색 경로에 적용된 규칙이지 전체 원장을 "
            "완전히 제한한 절대 규칙은 아니다."
        ),
        34: (
            "한 라운드는 최근 유전체에서 엘리트를 고른 뒤 교차·변이·무작위 탐색·전략 스펙으로 다음 "
            "개체군을 구성하면서 시작한다. 문법, 필드 가용성, 중복, 테마 조건을 사전 검사하고 통과한 "
            "후보만 시뮬레이션한다. 결과는 유전체, 부모 ID, 세대, 바뀐 유전자와 함께 저장된다. 사전 "
            "검사를 시뮬레이션 앞에 두면 이미 아는 실패를 줄일 수 있다. 동시 슬롯이 여덟 개여도 확실히 "
            "떨어질 후보를 태우면 다음 실험이 그만큼 밀린다."
        ),
        35: (
            "선택은 연속 적합도와 NSGA-II를 함께 쓴다[1]. 연속 적합도는 전 체크 통과 이전의 근접 "
            "실패에도 선택압을 준다. 본 연구가 전 체크 통과의 대리값으로 정의한 조건은 fail_count=0, "
            "Sharpe 기록, error_count=0이며 442건이 이를 만족했다. 전체의 1.2%다. NSGA-II는 Sharpe, "
            "Fitness, 제출 경로 점수, self-correlation, 최근 2년 Sharpe를 목적축으로 삼아 파레토 면과 "
            "혼잡 거리를 계산한다. Fitness에는 회전율 효율이 반영된다."
        ),
        46: (
            "Sharpe가 기록된 주요 자동 생성 경로 여섯 개(n=13,298)를 대상으로 공통 분석 기준 "
            "S≥1.58 도달 비율을 비교했다. 리전·지연·분류에 따라 실제 제출 조건은 달라질 수 있으므로 이 "
            "값을 모든 표본의 보편적 제출컷으로 해석하지 않는다. 전체 원장과 달리 기원이 없거나 다른 "
            "레거시 경로의 행도 표에서 제외했다."
        ),
        50: (
            "무작위 탐색은 2,582건 가운데 아홉 건이 분석 기준에 도달했다. 교차·스윕·개선·구제 경로는 "
            "17.4%에서 55.8%였다. 일반 변이는 부모를 쓰지만 수율은 8.6%였으므로 부모 유무만으로 차이를 "
            "설명할 수는 없다. 무작위 탐색은 신규 데이터셋에 처음 진입하는 역할이 있어 수율이 낮아도 "
            "탐색 폭을 위해 남겼다. 효율뿐 아니라 탐색 범위를 확보하기 위한 기준선이다[6]."
        ),
        68: (
            "생성 팔레트의 쏠림도 확인됐다. 자동 시딩 축 목록 26개가 모두 rsk70 데이터셋 소속이었고 "
            "2026년 7월 28일부터 8월 3일까지 alphas에서 submitted=1로 기록된 17건 가운데 11건이 그 "
            "데이터셋에서 나왔다. 다만 같은 기간 선택 규칙과 시딩 전략이 함께 바뀌었으므로 원인을 한 "
            "단계에만 돌릴 수는 없다."
        ),
        71: (
            "2026년 8월 13일부터 생성 우선순위를 피라미드 칸 기준으로 바꿨다. 칸은 리전·지연·데이터셋 "
            "카테고리의 조합이고 알파 세 건을 제출하면 하나의 피라미드로 집계된다[8]. 이미 세 건을 채운 "
            "칸에 더 제출해도 새 피라미드 수는 늘지 않는다. 그래서 미달 칸의 데이터셋을 먼저 뽑고 "
            "칸별 부족분만큼만 생성하도록 상한을 걸었다. 상한이 없으면 한 웨이브가 한 데이터셋으로 "
            "채워지는 현상을 운영 중 확인했기 때문이다."
        ),
        80: (
            "첫째, 평가자가 외부의 블랙박스다. 로컬에서 같은 결과를 재현할 수 없고 주간 테마와 일부 "
            "체크 조건이 달라질 수 있다. 시기별 성과를 단순 비교하면 교란이 생긴다."
        ),
        81: (
            "둘째, 시스템이 계속 개선됐다. 98일 동안 게이트 규칙, 선택 알고리즘, 시딩 전략이 여러 차례 "
            "바뀌었다. 기원별 수율과 변경 폭별 개선율은 운영 자료의 기술통계이며 통제 실험의 효과 "
            "추정치가 아니다. 시점 효과와 표본 구성 차이를 완전히 제거하지 못했다."
        ),
        85: (
            "유전 알고리즘은 자동으로 돈 버는 수식을 찾아 주는 장치가 아니다. 알파 실험을 구조화하는 "
            "장치다. 98일간의 운영이 남긴 것은 원장상 submitted=1인 86건에 그치지 않는다. 어떤 "
            "유전자가 어느 환경에서 작동했고 무엇이 실패했는지를 기록한 35,701건의 실험 원장이 더 "
            "중요한 결과다."
        ),
        87: (
            "제출 통과 건수를 운영 목표로 둔 기간에 원장은 submit_attempts의 submitted=1 기록 86건을 "
            "남겼다. 동시에 같은 계보를 반복해 제출하면서 프로덕션 상관이 높아졌다. 탐색 시스템의 "
            "목표는 자신이 만든 결과에 따라 제약 조건을 바꾸기도 한다. 이 되먹임을 기록하지 않으면 "
            "시스템은 이미 소진된 계보에 계속 예산을 쓸 수 있다."
        ),
        88: (
            "다음 단계로는 평가지표를 단순 제출 건수에서 프로덕션 상관 관문을 통과한 독립 계보 수로 "
            "바꾸는 작업을 두었다. 수량 전략이 스스로를 막는 구조를 지표 단계에서 줄이려는 것이다."
        ),
        90: (
            "[1] Deb, K., Pratap, A., Agarwal, S., & Meyarivan, T. (2002). A Fast and Elitist "
            "Multiobjective Genetic Algorithm: NSGA-II. IEEE Transactions on Evolutionary Computation, "
            "6(2), 182–197. https://doi.org/10.1109/4235.996017"
        ),
        93: (
            "[4] Thompson, W. R. (1933). On the Likelihood that One Unknown Probability Exceeds Another "
            "in View of the Evidence of Two Samples. Biometrika, 25(3/4), 285–294. "
            "https://doi.org/10.1093/biomet/25.3-4.285"
        ),
        94: (
            "[5] Chapelle, O., & Li, L. (2011). An Empirical Evaluation of Thompson Sampling. Advances in "
            "Neural Information Processing Systems, 24. "
            "https://papers.nips.cc/paper/2011/hash/e53a0a2978c28872a4505bdb51db06dc-Abstract.html"
        ),
        95: (
            "[6] Bergstra, J., & Bengio, Y. (2012). Random Search for Hyper-Parameter Optimization. Journal "
            "of Machine Learning Research, 13, 281–305. https://jmlr.org/papers/v13/bergstra12a.html"
        ),
        96: (
            "[7] Bailey, D. H., Borwein, J., López de Prado, M., & Zhu, Q. J. (2014). "
            "Pseudo-Mathematics and Financial Charlatanism: The Effects of Backtest Overfitting on "
            "Out-of-Sample Performance. Notices of the AMS, 61(5), 458–471. "
            "https://doi.org/10.1090/noti1105"
        ),
        97: (
            "[8] WorldQuant BRAIN. Consultant Dos and Don’ts; Getting Started: Power Pool Alphas; Getting "
            "Started: Finding Consultant Alphas; BRAIN Genius. Learn 문서, 2026-08-14 인출."
        ),
        101: (
            "본 보고서의 집계 시각은 2026년 8월 14일 16시 20분(KST)이며 대상 기간은 2026년 5월 "
            "9일부터 같은 날까지다. 전체 규모는 GenomicWQB 운영 데이터베이스 전체의 alphas 35,701행, "
            "rounds 923행, 고유 code 26,404종, 고유 settings_fp 2,227종이다. 제출 86건은 외부 "
            "리더보드가 아니라 submit_attempts에서 submitted=1인 기록을 센 값이다. 기원별 수율 표는 "
            "Sharpe가 기록된 주요 자동 생성 경로 여섯 개(n=13,298)만 포함한다. 전 체크 통과 442건은 "
            "본 연구가 정의한 대리값으로 Sharpe 기록, fail_count=0, error_count=0을 모두 만족한 행이다. "
            "계정 정보와 자격증명은 추출 대상에서 제외했다."
        ),
    }

    for index, text in replacements.items():
        set_paragraph(document.paragraphs[index], text)

    rq1 = document.tables[0].cell(1, 1)
    set_cell(
        rq1,
        "여러 유전자를 한꺼번에 바꾸는 변이보다 한 축 스윕과 실패 원인별 정향변이가 개선 효과를 "
        "더 분명하게 만드는가? 개선율과 후보의 기존 제출 풀 대비 자기상관을 함께 본다.",
    )

    genome = document.tables[1]
    set_cell(
        genome.cell(2, 2),
        "주요 창 2·3·5·10·20·60·120·252, 일부 경로 40",
    )
    set_cell(genome.cell(2, 3), "주요 창 우선·예외 잔존")

    appendix = document.tables[5]
    set_cell(
        appendix.cell(6, 2),
        "submit_attempts에서 submitted=1인 기록 86건과 거절 상태 문자열 파싱",
    )

    document.save(DOCX)


if __name__ == "__main__":
    main()
