#!/usr/bin/env python3
"""Render the GenomicWQB rearchitecture proposal from Markdown to DOCX.

The document contains project research results but no personal identity data, and
is written directly to the vault-backed document directory supplied by caller.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


FONT = "Noto Sans CJK KR"
NAVY = RGBColor(20, 37, 63)
TEAL = RGBColor(0, 121, 140)
INK = RGBColor(35, 42, 52)
GRAY = RGBColor(95, 105, 115)
PALE = "EEF3F6"
LIGHT_TEAL = "E8F4F5"
WHITE = "FFFFFF"


def args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("source", type=Path)
    p.add_argument("output", type=Path)
    return p.parse_args()


def font(run, size=None, bold=None, color=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def no_row_split(row, header=False):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    if header:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.extend((begin, instr, separate, end))
    font(r, size=8, color=GRAY)


def set_repeat_table_header(row):
    no_row_split(row, header=True)


def set_cell_text(cell, text, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    text = re.sub(r"`([^`]*)`", r"\1", text).replace("**", "")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.strip())
    font(r, size=8.6, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(21)
    section.bottom_margin = Mm(19)
    section.left_margin = Mm(23)
    section.right_margin = Mm(23)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.42
    normal.paragraph_format.space_after = Pt(5.2)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (25, NAVY, 26, 12),
        "Heading 2": (16, NAVY, 18, 8),
        "Heading 3": (11.5, TEAL, 12, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        s = styles[name]
        s.font.name = FONT
        s._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    if "Proposal Subtitle" not in styles:
        subtitle = styles.add_style("Proposal Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Proposal Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    subtitle.font.size = Pt(14)
    subtitle.font.bold = True
    subtitle.font.color.rgb = TEAL
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    if "Small Meta" not in styles:
        meta = styles.add_style("Small Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Small Meta"]
    meta.font.name = FONT
    meta._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    meta.font.size = Pt(9)
    meta.font.color.rgb = GRAY
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.text = "GENOMICWQB 2.0  ·  REARCHITECTURE PROPOSAL"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.runs[0], size=7.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "GenomicWQB 전면 재구성 제안서"
    props.subject = "증거 기반 포트폴리오 알파 탐색 시스템"
    props.author = "GenomicWQB"
    props.keywords = "GenomicWQB, genetic algorithm, alpha research, AutoZoom, evidence ledger"


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, base_size=9.5, base_color=INK):
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = paragraph.add_run(piece[2:-2])
            font(r, size=base_size, bold=True, color=NAVY)
        elif piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])
            font(r, size=base_size - 0.4, color=TEAL, name="Noto Sans Mono CJK KR")
        else:
            r = paragraph.add_run(piece)
            font(r, size=base_size, color=base_color)


def add_quote(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.7)
    cell = table.cell(0, 0)
    shade(cell, LIGHT_TEAL)
    cell_margins(cell, top=170, start=220, bottom=170, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, base_size=11, base_color=NAVY)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    widths = None
    if cols == 4 and rows[0][0].strip() == "번호":
        widths = [12, 70, 12, 70]
        table.autofit = False
        for j, width in enumerate(widths):
            table.columns[j].width = Mm(width)
    for i, row in enumerate(rows):
        no_row_split(table.rows[i], header=(i == 0))
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            if widths:
                cell.width = Mm(widths[j])
            if i == 0:
                shade(cell, "14253F")
                set_cell_text(cell, text, bold=True, color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if i % 2 == 0:
                    shade(cell, PALE)
                numeric = bool(re.fullmatch(r"[\d,.% /~+\-]+", text.strip()))
                set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER if numeric else WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, number: str | None = None):
    p = doc.add_paragraph(style=None if number is not None else "List Bullet")
    p.paragraph_format.left_indent = Mm(7)
    p.paragraph_format.first_line_indent = Mm(-3.5)
    p.paragraph_format.space_after = Pt(3.5)
    if number is not None:
        prefix = p.add_run(f"{number}.  ")
        font(prefix, bold=True, color=TEAL)
    add_inline(p, text)


def render(source: Path, output: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure(doc)
    title_seen = False
    subtitle_seen = False
    paragraph_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_paragraph():
        nonlocal paragraph_buf
        if not paragraph_buf:
            return
        text = " ".join(x.strip() for x in paragraph_buf).strip()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Mm(4)
        add_inline(p, text)
        paragraph_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            filtered = [r for r in table_buf if not all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in r)]
            add_table(doc, filtered)
            table_buf = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("<!-- HUMANIZE-SUMMARY"):
            break
        if table_buf and not line.startswith("|"):
            flush_table()
        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            table_buf.append([c.strip() for c in line.strip("|").split("|")])
            continue
        if not line.strip():
            flush_paragraph()
            continue
        if line.strip() == "---":
            flush_paragraph()
            flush_table()
            doc.add_page_break()
            continue
        if line.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(74)
            add_inline(p, line[2:].strip(), base_size=25, base_color=NAVY)
            title_seen = True
            continue
        if line.startswith("## "):
            flush_paragraph()
            text = line[3:].strip()
            if title_seen and not subtitle_seen and text.startswith("제출량 중심"):
                p = doc.add_paragraph(style="Proposal Subtitle")
                add_inline(p, text, base_size=14, base_color=TEAL)
                subtitle_seen = True
            else:
                p = doc.add_paragraph(style="Heading 1")
                add_inline(p, text, base_size=16, base_color=NAVY)
            continue
        if line.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip(), base_size=11.5, base_color=TEAL)
            continue
        if re.match(r"^작성 기준일:|^대상:|^문서 성격:", line):
            flush_paragraph()
            p = doc.add_paragraph(style="Small Meta")
            add_inline(p, line.replace("  ", ""), base_size=9, base_color=GRAY)
            continue
        if line.startswith("> "):
            flush_paragraph()
            add_quote(doc, line[2:].strip())
            continue
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            flush_paragraph()
            add_bullet(doc, m.group(2), number=m.group(1))
            continue
        if line.startswith("- "):
            flush_paragraph()
            add_bullet(doc, line[2:])
            continue
        if line.endswith("  "):
            paragraph_buf.append(line[:-2])
            flush_paragraph()
            continue
        paragraph_buf.append(line)

    flush_paragraph()
    flush_table()

    # Avoid a trailing blank page created by the last explicit break.
    body = doc._element.body
    paragraphs = body.findall(qn("w:p"))
    if paragraphs:
        last = paragraphs[-1]
        br = last.find(".//" + qn("w:br"))
        if br is not None and br.get(qn("w:type")) == "page":
            body.remove(last)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    ns = args()
    render(ns.source, ns.output)
