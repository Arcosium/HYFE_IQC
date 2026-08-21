#!/usr/bin/env python3
"""Build the renewed single-hypothesis alpha-research report.

The output document is written to the vault-backed target supplied by the
caller. Author metadata is preserved by reading it from the existing report;
no identity value is embedded in this repository.
"""

from __future__ import annotations

import argparse
import json
import math
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


FONT = "Noto Serif CJK KR"
SERIF = "Noto Serif CJK KR"
NAVY = RGBColor(0, 0, 0)
BLUE = RGBColor(0, 0, 0)
ORANGE = RGBColor(0, 0, 0)
GRAY = RGBColor(70, 70, 70)
LIGHT = "E7E7E7"
PALE_ORANGE = "F2F2F2"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--existing", type=Path, required=True)
    parser.add_argument("--analysis", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--source-text", type=Path, required=True)
    return parser.parse_args()


def set_run_font(run, name=FONT, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=85, bottom=70, end=85) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    set_run_font(run, size=8, color=GRAY)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "245C88")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    new_run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.2)
    normal.paragraph_format.line_spacing = 1.62
    normal.paragraph_format.space_after = Pt(4.8)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = SERIF
    title._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), SERIF)
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title_properties = title._element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)

    for name, size, color in (
        ("Heading 1", 14.5, RGBColor(0, 0, 0)),
        ("Heading 2", 11.5, RGBColor(0, 0, 0)),
        ("Heading 3", 10.5, RGBColor(0, 0, 0)),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 9)
        style.paragraph_format.space_after = Pt(6)

    if "Report Subtitle" not in styles:
        subtitle = styles.add_style("Report Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Report Subtitle"]
    subtitle.font.name = SERIF
    subtitle._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), SERIF)
    subtitle.font.size = Pt(12.5)
    subtitle.font.bold = True
    subtitle.font.color.rgb = RGBColor(0, 0, 0)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)

    if "Abstract" not in styles:
        abstract = styles.add_style("Abstract", WD_STYLE_TYPE.PARAGRAPH)
    else:
        abstract = styles["Abstract"]
    abstract.font.name = FONT
    abstract._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    abstract.font.size = Pt(9.4)
    abstract.paragraph_format.line_spacing = 1.5
    abstract.paragraph_format.space_after = Pt(4)
    abstract.paragraph_format.first_line_indent = Mm(4)

    if "Caption Custom" not in styles:
        caption = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption Custom"]
    caption.font.name = FONT
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(3)

    if "Reference" not in styles:
        reference = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = styles["Reference"]
    reference.font.name = FONT
    reference._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    reference.font.size = Pt(8.4)
    reference.paragraph_format.left_indent = Mm(6)
    reference.paragraph_format.first_line_indent = Mm(-6)
    reference.paragraph_format.line_spacing = 1.15
    reference.paragraph_format.space_after = Pt(2.2)


def add_body(doc: Document, text: str, style: str | None = None, indent=True):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    if indent and style is None:
        paragraph.paragraph_format.first_line_indent = Mm(4)
    return paragraph


def add_note(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_ORANGE)
    set_cell_margins(cell, top=105, start=140, bottom=105, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.9, color=RGBColor(0, 0, 0))
    return table


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    repeat_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        shade_cell(cell, "D9D9D9")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=7.8, bold=True, color=RGBColor(0, 0, 0))
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 1:
                shade_cell(cell, "F2F2F2")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(text))
            set_run_font(run, size=7.6)
            if idx > 0 and re.fullmatch(r"[-+−]?[\d,.%()]+", str(text).strip()):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if widths:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Mm(width)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    return table


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(text, style="Caption Custom")
    return paragraph


def add_figure(doc: Document, path: Path, caption: str, width_mm: float):
    add_caption(doc, caption)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(path), width=Mm(width_mm))
    return paragraph


def percent(value: float | None, digits=1) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "NA"
    return f"{value * 100:.{digits}f}%"


def number(value: float | None, digits=3, sign=False) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "NA"
    prefix = "+" if sign and value > 0 else ""
    return f"{prefix}{value:.{digits}f}"


def ci(low: float, high: float, scale=1.0, digits=2) -> str:
    return f"[{low * scale:.{digits}f}, {high * scale:.{digits}f}]"


def preserve_identity(existing: Path) -> str:
    source = Document(existing)
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("저자:"):
            return text
    return "저자 정보는 기존 문서와 동일"


def start_section(doc: Document, page_start: int | None = None):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    if page_start is not None:
        properties = section._sectPr
        page_number = properties.find(qn("w:pgNumType"))
        if page_number is None:
            page_number = OxmlElement("w:pgNumType")
            properties.append(page_number)
        page_number.set(qn("w:start"), str(page_start))
    return section


def add_front_matter(doc: Document, identity: str, summary: dict) -> None:
    snapshot = summary["snapshot"]
    effects = summary["effects"]
    regressions = summary["regressions"]
    external = summary["external"]

    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)
    doc.add_paragraph("유전 알고리즘으로 알파 리서치하기", style="Title")
    doc.add_paragraph("1축 변이는 다축 변이보다 효과적인가", style="Report Subtitle")
    doc.add_paragraph(
        "공개 최적화 문제의 통제 실험과 GenomicWQB 퀀트 알파 계보의 실증 분석",
        style="Report Subtitle",
    )
    english = doc.add_paragraph()
    english.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = english.add_run(
        "Is Single-Axis Mutation Effective in Evolutionary Alpha Search?\n"
        "From Controlled IOHprofiler Problems to GenomicWQB Lineages"
    )
    set_run_font(run, name="Liberation Serif", size=10.5, italic=True, color=GRAY)
    english.paragraph_format.space_after = Pt(34)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(author.add_run(identity), size=10, color=RGBColor(0, 0, 0))
    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        date_line.add_run(
            "2026년 8월 21일"
        ),
        size=9.5,
        color=RGBColor(0, 0, 0),
    )
    start_section(doc)

    abstract_heading = doc.add_paragraph()
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.paragraph_format.space_after = Pt(12)
    set_run_font(abstract_heading.add_run("국문 초록"), size=15, bold=True)
    abstract = (
        "본 연구는 유전 알고리즘이 좋은 후보 주변을 다듬는 단계에서 한 번에 축 하나만 바꾸는 1축 변이가 유효한지 검정한다. "
        "먼저 변이 폭만 통제할 수 있는 IOHprofiler PBO 공개 문제로 작동 원리를 확인했다. 25개 문제를 크기 2종과 문제별 설정 3종으로 나누고 각각 20회 반복해 "
        f"{external['paired_runs']:,}쌍을 비교했다. 1축은 최적점에 도달한 비율이 {percent(external['hit_single'], 1)}로 다축의 {percent(external['hit_multi'], 1)}보다 높았다. "
        f"평균 이동 거리는 다축이 조금 앞섰으나 차이의 95% 신뢰구간은 0을 포함했다(1축−다축 {number(external['progress_difference'], 3, sign=True)}, {ci(external['progress_ci_low'], external['progress_ci_high'], 1, 3)}). "
        "기본 문제에서는 1축이 빠르게 최적점에 모였고 여러 축을 함께 바꿔야 풀리는 문제에서는 다축이 멀리 이동했다. 공개 문제 결과는 1축을 모든 상황의 우월한 방식이 아니라 국소 탐색에 맞는 방식으로 설명한다. "
        f"이 원리가 퀀트 알파에서도 나타나는지 GenomicWQB 기록 {snapshot['alphas']:,}건으로 다시 확인했다. 캐시와 무변경, 완전 동일 설정을 제외한 실제 부모·자식은 {summary['cohort_flow']['primary_evaluated_lineage_pairs']:,}쌍이었다. "
        f"부모보다 Sharpe가 좋아진 비율은 1축 {percent(effects['improved']['single'], 2)}, 다축 {percent(effects['improved']['multi'], 2)}로 {effects['improved']['difference'] * 100:.2f}퍼센트포인트 차이가 났다(95% CI {ci(effects['improved']['ci_low'], effects['improved']['ci_high'], 100, 2)}). "
        f"운영 조건을 맞춘 차이는 {regressions['adjusted_improved']['estimate'] * 100:.2f}퍼센트포인트였고 같은 부모의 자식끼리 비교해도 {regressions['sibling_fe_improved']['estimate'] * 100:.2f}퍼센트포인트였다. "
        "다만 2축도 1축과 비슷했고 Sharpe 1.58을 새로 넘는 비율은 늘지 않았다. 서로 다른 두 자료가 같은 방향을 보였다는 점에서 1축 변이는 퀀트 알파의 국소 탐색에서도 유효하다. 운용할 때는 좋은 부모를 다듬는 단계에서 1~2축을 우선하고 개선이 멈추면 넓은 변이를 섞는 편이 맞다."
    )
    add_body(doc, abstract, style="Abstract", indent=False)
    keywords = doc.add_paragraph()
    set_run_font(
        keywords.add_run(
            "주제어: 유전 알고리즘, 알파 리서치, 1축 변이, 국소 탐색, IOHprofiler, PBO"
        ),
        size=9,
        bold=True,
        color=RGBColor(0, 0, 0),
    )
    jel = doc.add_paragraph()
    set_run_font(jel.add_run("JEL 분류: C61, C63, G17"), size=9, italic=True)

    doc.add_page_break()
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = Pt(18)
    set_run_font(toc_heading.add_run("목차"), size=15, bold=True)
    contents = [
        "1. 서론",
        "2. 이론적 배경과 연구가설",
        "3. IOHprofiler 공개 문제와 통제 실험",
        "4. 외부 실험 결과",
        "5. GenomicWQB 자료와 분석 방법",
        "6. GenomicWQB 실증 결과",
        "7. 강건성 검토와 증거 통합",
        "8. 운용 제안과 후속 검증",
        "9. 연구의 한계",
        "10. 결론",
        "11. 참고문헌",
        "부록",
    ]
    for item in contents:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(18)
        paragraph.paragraph_format.space_after = Pt(7)
        set_run_font(paragraph.add_run(item), size=10.5, bold=item.split(".", 1)[0].isdigit())
    start_section(doc, page_start=1)


def build_report(doc: Document, analysis_dir: Path, summary: dict) -> None:
    snapshot = summary["snapshot"]
    flow = summary["cohort_flow"]
    effects = summary["effects"]
    regressions = summary["regressions"]
    external = summary["external"]
    by_k = pd.read_csv(analysis_dir / "internal_by_mutation_width.csv")
    hetero = pd.read_csv(analysis_dir / "internal_heterogeneity.csv")
    genes = pd.read_csv(analysis_dir / "internal_single_gene_results.csv")
    cells = pd.read_csv(analysis_dir / "external_benchmark_cells.csv")

    doc.add_heading("1. 문제를 한 문장으로 줄이기", level=1)
    add_body(
        doc,
        "알파 후보 하나에는 데이터필드, 변환 연산자, 룩백, 중립화, 감쇠, 유니버스 같은 여러 설정이 들어간다. 유전 알고리즘은 기존 알파의 설정 일부를 바꿔 자식을 만들고 시뮬레이션으로 성적을 확인한다. 평가비용이 작다면 여러 설정을 크게 바꾸고 많이 시험해도 된다. WorldQuant BRAIN처럼 결과를 확인하는 데 시간과 실행 슬롯이 필요하면 이야기가 달라진다. 한 번에 몇 개를 바꿀지가 곧 예산 문제다.",
    )
    add_body(
        doc,
        "기존 리포트는 두 가지를 물었다. 하나는 몇 개의 축을 바꿀지였고 다른 하나는 여러 성과 기준을 함께 쓰는 선택 방식이었다. 이번에는 두 번째 질문을 뺐다. 여러 축을 한꺼번에 바꾸는 것보다 한 축만 바꾸는 편이 나은가만 살펴본다. 대신 최신 기록까지 표본을 늘리고 같은 부모에서 나온 자식끼리 비교했다. 공개 문제를 이용한 별도 실험도 더했다.",
    )
    add_note(
        doc,
        "여기서 1축은 genes_changed에 기록된 유전자가 정확히 하나라는 뜻이다. 축의 개수는 변화의 수를 세며 변화량의 크기나 수식 의미의 거리를 재지 않는다.",
    )

    doc.add_heading("1.1 적게 바꾸는 방법과 많이 바꾸는 방법", level=2)
    add_body(
        doc,
        "1축 변이는 부모 구조를 거의 그대로 둔다. 무엇을 바꿔서 성적이 달라졌는지 찾기 쉽고 좋은 부모 주변을 꼼꼼히 살필 수 있다. 여러 축을 함께 바꾸면 한 번에 더 멀리 이동한다. 가까운 곳에서 더 좋아질 길이 막혔거나 비슷한 알파만 계속 나올 때는 이 방법이 도움이 될 수 있다. 어느 쪽이 좋은지는 문제의 모양과 부모의 위치, 축끼리 영향을 주고받는 정도에 달려 있다. OneMax처럼 한 축씩 고쳐도 되는 문제에서는 작은 변이가 강하다. 여러 축을 함께 바꿔야 벽을 넘는 문제에서는 큰 변이가 더 빠를 수 있다[5,6].",
    )
    add_body(
        doc,
        "알파에는 성격이 다른 축이 섞여 있다. 룩백이나 감쇠는 가까운 값을 차례로 비교하기 쉽다. 데이터필드나 결합 방식은 조금만 바꿔도 알파의 경제적 의미가 크게 달라질 수 있다. 운영 기록만으로 이런 차이를 모두 알 수는 없다. 그래서 실제 운영 기록과 조건을 맞춘 공개 문제 실험을 따로 분석했다.",
    )

    doc.add_heading("2. 하나의 가설과 확인 기준", level=1)
    add_body(
        doc,
        "연구가설 H1은 같은 탐색 단계라면 1축 변이가 다축 변이보다 부모의 Sharpe를 높일 가능성이 크다는 것이다. 반대로 차이가 없거나 다축이 더 낫다고 보는 기준을 H0으로 둔다. 가장 중요한 지표는 자식 Sharpe가 부모보다 높아졌는지다. Sharpe가 얼마나 달라졌는지와 1.58을 새로 넘었는지도 함께 본다. 기존 제출 알파와 얼마나 비슷한지는 self-correlation으로 확인한다.",
    )
    add_body(
        doc,
        "효과를 숫자 하나로만 판단하지 않았다. 부모보다 조금 나아지는 사례가 많아도 제출에 가까운 기준을 넘지 못할 수 있다. 평균 하락폭이 작다고 해서 최적점에 더 자주 닿는 것도 아니다. 지표마다 답이 다르면 그 차이까지 결과로 받아들였다.",
    )
    design_rows = [
        ["처치", "genes_changed 길이=1"],
        ["비교", "genes_changed 길이≥2"],
        ["가장 중요한 지표", "자식 Sharpe > 부모 Sharpe"],
        ["함께 볼 지표", "ΔSharpe, S≥1.58 신규 도달, self-correlation"],
        ["내부 비교 단위", "캐시가 아닌 부모·자식 실평가 한 쌍"],
        ["외부 비교 단위", "같은 문제·시작점·평가예산을 쓴 실행 한 쌍"],
        ["오차 범위", "부모 묶음 95% CI, 문제×크기 묶음 95% CI"],
    ]
    add_caption(doc, "표 1. 무엇을 어떻게 비교했는가")
    add_table(doc, ["항목", "정의"], design_rows, widths=[35, 125])

    doc.add_heading("3. 자료", level=1)
    doc.add_heading("3.1 GenomicWQB 운영 원장", level=2)
    add_body(
        doc,
        f"내부 자료는 {snapshot['max_kst'].replace('T', ' ')} 시점의 SQLite 복사본이다. 기록 기간은 {snapshot['min_kst'][:10]}부터 {snapshot['max_kst'][:10]}까지다. 알파 {snapshot['alphas']:,}건과 라운드 {snapshot['rounds']:,}건이 들어 있다. 서로 다른 표현식은 {snapshot['unique_codes']:,}종, 서로 다른 설정은 {snapshot['unique_settings']:,}종이다. 제출 성공 기록은 {snapshot['successful_submit_attempts']:,}건이지만 제출 수는 이번 가설을 판단하는 지표로 쓰지 않았다.",
    )
    add_body(
        doc,
        "부모와 자식의 Sharpe가 모두 있는 기록을 연결하고 genes_changed에서 바꾼 축 수를 셌다. 아무것도 바꾸지 않은 기록은 뺐다. 캐시를 다시 쓴 사례는 새 평가를 돌리지 않았으므로 제외했다. 표현식과 설정이 모두 같은 사례도 변이 효과로 보기 어려워 뺐다. 같은 부모에서 나온 자식은 서로 닮을 수 있다. 계산할 때 이들을 부모 ID별로 한 묶음으로 다뤘다[8].",
    )
    flow_rows = [
        ["부모·자식 Sharpe 연결", f"{flow['joined_parent_child_with_sharpe']:,}", "시작 표본"],
        ["변경축 0 제외", f"−{flow['excluded_zero_changed']:,}", "무변경·기록 불일치"],
        ["캐시 재사용 제외", f"−{flow['excluded_cached']:,}", "신규 평가비용 없음"],
        ["완전 동일 설정 제외", f"−{flow['excluded_exact_config']:,}", "code_hash와 settings_fp 모두 동일"],
        ["시간순서 오류 제외", f"−{flow['excluded_invalid_time_order']:,}", "자식 시각이 부모보다 이른 경우"],
        ["주 분석 표본", f"{flow['primary_evaluated_lineage_pairs']:,}", f"1축 {flow['primary_single']:,} · 다축 {flow['primary_multi']:,}"],
    ]
    add_caption(doc, "표 2. 내부 분석 대상을 고른 과정")
    add_table(doc, ["단계", "행 수", "판정"], flow_rows, widths=[52, 28, 80])

    doc.add_heading("3.2 IOHprofiler 공개 문제와 재실험 자료", level=2)
    add_body(
        doc,
        "외부 실험에는 IOHprofiler의 PBO 문제 25종을 썼다. PBO는 0과 1로 된 답을 조금씩 바꾸며 가장 좋은 값을 찾는 공개 문제 모음이다. OneMax, LeadingOnes, W-model 변형, Ising, N-Queens, NK Landscape 등이 들어 있다[3,4]. 실제 알파와 생김새는 다르지만 한 번에 몇 개를 바꿀지 정확히 맞출 수 있다. 변이 폭 자체의 효과를 따로 확인하기에 알맞다.",
    )
    add_body(
        doc,
        f"각 문제는 크기 16과 100으로 나눴고 문제별 설정 1~3을 사용했다. 조합마다 시작점 20개를 만들었다. 1축과 다축은 같은 시작점에서 출발하고 똑같이 20d번 평가받았다. 전체 비교는 {external['paired_runs']:,}쌍이다. 두 알고리즘을 합치면 {external['paired_runs'] * 2:,}회 실행했고 값을 계산한 횟수는 6,960,000회다. 다축은 GenomicWQB에서 실제로 나타난 2~14축 비율을 그대로 따랐다. 바꾸는 축 수는 내부 기록과 연결하되 결과는 별도의 공개 문제에서 얻었다.",
    )
    sources_rows = [
        ["GenomicWQB", "실제 운영 기록", f"알파 {snapshot['alphas']:,}건", "실제 운용 결과 확인"],
        ["IOH PBO", "공개 표준 문제", "25문제×2크기×3설정", "변이 폭만 바꿔 비교"],
        ["IOHdata", "공개 원본 ZIP", "RLS/self_GA 예제", "자료 형식과 구현 확인"],
        ["문헌", "이론·비교 연구", "작은 변이와 큰 점프의 경계", "결과 해석"],
    ]
    add_caption(doc, "표 3. 자료원의 역할")
    add_table(doc, ["자료원", "성격", "범위", "분석 역할"], sources_rows, widths=[34, 35, 45, 50])

    doc.add_heading("4. 비교 방법", level=1)
    doc.add_heading("4.1 바꾼 축 수별 결과 보기", level=2)
    add_body(
        doc,
        "바꾼 축 수를 1, 2, 3, 4, 5, 6개 이상으로 나눴다. 각 구간에서 부모보다 좋아진 비율과 95% 신뢰구간을 계산했다. 가설 판단에는 1축과 다축의 비교만 썼다. 축 수별 세부 결과와 어떤 유전자를 바꿨는지는 참고용으로 따로 봤다. 축을 많이 바꿀수록 결과가 계속 나빠지는지, 2축과 3축 사이에서 크게 달라지는지도 확인했다.",
    )
    doc.add_heading("4.2 같은 부모의 자식을 한 묶음으로 계산", level=2)
    add_body(
        doc,
        "같은 부모에서 나온 자식은 출발 성적과 실패 원인이 비슷하다. 모든 자식을 서로 무관한 사례로 보면 오차 범위가 실제보다 좁아질 수 있다. 이를 막으려고 부모를 한 묶음으로 삼아 2,000번 다시 뽑았다. 이 방법을 부모 단위 군집 부트스트랩이라고 한다. 조건을 맞추는 통계 계산에서도 부모 ID가 같은 자식을 한 묶음으로 처리했다[8].",
    )
    add_body(
        doc,
        "운영 조건이 달라서 생길 수 있는 차이도 줄였다. 부모 Sharpe, 생성 세대와 방식, 리전, 유니버스, 지연, 평가 주차를 통계적으로 맞췄다. 계산 결과는 조건을 맞춘 뒤 1축과 다축의 개선률이 몇 퍼센트포인트 다른지를 뜻한다. 다만 기록에 없는 차이까지 없앨 수는 없다. 실제 무작위 실험을 대신하는 분석은 아니다.",
    )
    doc.add_heading("4.3 같은 부모에서 나온 자식끼리 비교", level=2)
    add_body(
        doc,
        f"1축 자식과 다축 자식을 모두 만든 부모는 {regressions['mixed_parent_count']:,}개였고 해당 자식은 {regressions['mixed_parent_rows']:,}건이었다. 이 자식끼리 비교하면 부모 자체의 품질 차이를 줄일 수 있다. 생성 방식과 평가 주차도 맞췄다. 자식 수가 많은 부모가 결과를 좌우하지 않도록 부모마다 같은 무게를 주는 계산도 덧붙였다. 그래도 운영 시스템이 왜 어떤 부모에는 1축을, 다른 부모에는 다축을 골랐는지까지 알 수는 없다.",
    )
    doc.add_heading("4.4 외부 실험의 조건을 똑같이 맞추기", level=2)
    add_body(
        doc,
        "두 알고리즘은 같은 규칙을 썼다. 새 후보가 현재 후보보다 좋거나 같을 때만 다음 부모가 됐다. 1축은 매번 한 자리만 뒤집었다. 다축은 GenomicWQB에서 나타난 비율에 맞춰 k를 뽑고 k자리를 뒤집었다. 결과는 두 가지로 쟀다. 하나는 정해진 횟수 안에 최적점을 찾았는지다. 다른 하나는 시작점과 최적점 사이의 거리를 얼마나 줄였는지다. 뒤의 값을 정규화 진척도라고 부른다. 오차 범위는 문제와 크기가 같은 50개 묶음을 기준으로 계산했다.",
    )
    add_note(
        doc,
        "공개 문제 실험은 알파 수익성을 재현하지 않는다. 여러 축을 함께 바꿔야 풀리는 문제인지에 따라 1축과 다축의 우위가 달라지는지만 확인한다.",
    )

    doc.add_heading("5. GenomicWQB 결과", level=1)
    doc.add_heading("5.1 폭이 세 축을 넘으면 개선률이 급격히 낮아졌다", level=2)
    by_k_rows = []
    for row in by_k.itertuples():
        by_k_rows.append(
            [
                f"{row.k_group}축",
                f"{int(row.n):,}",
                percent(row.improve_rate, 2),
                ci(row.improve_ci_low, row.improve_ci_high, 100, 2),
                number(row.delta_median, 2, sign=True),
                percent(row.crossed_158_rate, 2),
            ]
        )
    add_caption(doc, "표 4. 변경 유전자 수별 내부 성과")
    add_table(
        doc,
        ["변이 폭", "n", "개선률", "95% CI", "ΔSharpe 중앙값", "S≥1.58 신규 도달"],
        by_k_rows,
        widths=[24, 24, 27, 35, 33, 34],
    )
    add_figure(
        doc,
        analysis_dir / "fig1_internal_dose_response.png",
        "그림 1. 변경 유전자 수와 부모 대비 개선률·Sharpe 변화량",
        164,
    )
    add_body(
        doc,
        f"1축에서 부모보다 좋아진 비율은 {percent(by_k.iloc[0]['improve_rate'], 2)}였다. 2축은 {percent(by_k.iloc[1]['improve_rate'], 2)}로 조금 높았지만 오차 범위가 크게 겹쳤다. 3축은 {percent(by_k.iloc[2]['improve_rate'], 2)}, 4축은 {percent(by_k.iloc[3]['improve_rate'], 2)}로 내려갔다. 5축과 6축 이상은 약 6%였다. 자료는 1축만 특별히 좋다고 말하지 않는다. 1~2축처럼 적게 바꿀 때와 3축 이상을 넓게 바꿀 때의 차이가 더 뚜렷하다.",
    )
    add_body(
        doc,
        "어느 구간이든 Sharpe 변화량의 중앙값은 0보다 작았다. 새로 만든 자식은 대체로 부모보다 나빴다. 1축이 평균 성적을 올렸다는 뜻은 아니다. 다축보다 실패가 적고 떨어지는 폭도 작았다는 뜻이다. 따라서 개선률 16%만 보고 1축이 대체로 성공한다고 이해하면 안 된다.",
    )

    doc.add_heading("5.2 좋은 부모를 조금 다듬을 때는 1축이 나았다", level=2)
    effect_rows = [
        [
            "단순 비교",
            f"{effects['improved']['difference'] * 100:+.2f}pp",
            ci(effects['improved']['ci_low'], effects['improved']['ci_high'], 100, 2),
            f"n={effects['improved']['single_n'] + effects['improved']['multi_n']:,}",
        ],
        [
            "운영 조건을 맞춘 비교",
            f"{regressions['adjusted_improved']['estimate'] * 100:+.2f}pp",
            ci(regressions['adjusted_improved']['ci_low'], regressions['adjusted_improved']['ci_high'], 100, 2),
            f"p={regressions['adjusted_improved']['p_value']:.4g}",
        ],
        [
            "같은 부모의 자식 비교",
            f"{regressions['sibling_fe_improved']['estimate'] * 100:+.2f}pp",
            ci(regressions['sibling_fe_improved']['ci_low'], regressions['sibling_fe_improved']['ci_high'], 100, 2),
            f"부모 {regressions['mixed_parent_count']:,}개",
        ],
        [
            "부모마다 같은 무게",
            f"{regressions['sibling_equal_parent_improved']['estimate'] * 100:+.2f}pp",
            ci(regressions['sibling_equal_parent_improved']['ci_low'], regressions['sibling_equal_parent_improved']['ci_high'], 100, 2),
            "부모당 같은 가중치",
        ],
    ]
    add_caption(doc, "표 5. 1축 변이의 개선률 차이")
    add_table(doc, ["추정", "효과", "95% CI", "표본·검정"], effect_rows, widths=[48, 28, 42, 48])
    add_figure(
        doc,
        analysis_dir / "fig2_internal_effect_estimates.png",
        "그림 2. 계산 방법별 1축 변이의 개선률 차이",
        133,
    )
    add_body(
        doc,
        f"그대로 비교하면 1축의 개선률이 {effects['improved']['difference'] * 100:.2f}퍼센트포인트 높았다. 운영 조건을 맞추면 차이는 {regressions['adjusted_improved']['estimate'] * 100:.2f}퍼센트포인트였다. 같은 부모에서 나온 자식끼리 비교해도 {regressions['sibling_fe_improved']['estimate'] * 100:.2f}퍼센트포인트 높았다. 세 계산에서 방향과 크기가 비슷했다. 기록으로 확인할 수 있는 부모 품질과 평가 시점만으로는 이 차이를 설명하기 어렵다.",
    )
    add_body(
        doc,
        f"Sharpe가 얼마나 달라졌는지도 비슷했다. 단순 평균 차이는 {effects['delta_sharpe']['difference']:+.3f}였다. 운영 조건을 맞춘 차이는 {regressions['adjusted_delta_sharpe']['estimate']:+.3f}(95% CI {ci(regressions['adjusted_delta_sharpe']['ci_low'], regressions['adjusted_delta_sharpe']['ci_high'], 1, 3)}), 같은 부모의 자식끼리 비교한 차이는 {regressions['sibling_fe_delta_sharpe']['estimate']:+.3f}이었다. 1축 자식도 평균적으로 부모보다 나빴다. 다만 다축 자식보다 덜 떨어졌다.",
    )

    doc.add_heading("5.3 최종 문턱 도달은 개선되지 않았다", level=2)
    add_body(
        doc,
        f"Sharpe 1.58을 새로 넘은 비율은 1축 {percent(effects['crossed_158']['single'], 2)}, 다축 {percent(effects['crossed_158']['multi'], 2)}였다. 단순 차이는 {effects['crossed_158']['difference'] * 100:+.2f}퍼센트포인트였고 95% 신뢰구간은 {ci(effects['crossed_158']['ci_low'], effects['crossed_158']['ci_high'], 100, 2)}였다. 운영 조건을 맞춰 계산해도 차이는 {regressions['adjusted_crossed_158']['estimate'] * 100:+.2f}퍼센트포인트에 그쳤다. 1축이 낫다고 보기 어려운 결과다.",
    )
    add_body(
        doc,
        f"같은 부모의 자식끼리 비교하면 차이는 {regressions['sibling_fe_crossed_158']['estimate'] * 100:+.2f}퍼센트포인트였다. 적게 바꾸면 작은 개선은 자주 만들 수 있다. 하지만 성적이 낮은 부모를 한 번에 기준 밖으로 끌어올리는 힘은 약할 수 있다. 부모보다 조금 좋아질 가능성은 높았지만 최종 기준을 넘는 비율까지 좋아지지는 않았다.",
    )

    doc.add_heading("5.4 자식을 만든 방식과 부모 성적에 따라 달랐다", level=2)
    origin = hetero[hetero["dimension"] == "origin"].copy()
    quartile = hetero[hetero["dimension"] == "parent_quartile"].copy()
    hetero_rows = []
    for row in pd.concat([origin, quartile]).itertuples():
        label = {"crossover": "교차", "mutate": "일반 변이", "sweep": "스윕"}.get(row.group, row.group)
        hetero_rows.append(
            [
                "생성 경로" if row.dimension == "origin" else "부모 Sharpe",
                label,
                f"{int(row.n):,}",
                percent(row.single_rate, 1),
                percent(row.multi_rate, 1),
                f"{row.difference * 100:+.1f}pp",
                ci(row.ci_low, row.ci_high, 100, 1),
            ]
        )
    add_caption(doc, "표 6. 하위집단별 개선률 차이")
    add_table(doc, ["구분", "집단", "n", "1축", "다축", "차이", "95% CI"], hetero_rows)
    add_figure(
        doc,
        analysis_dir / "fig3_internal_heterogeneity.png",
        "그림 3. 자식을 만든 방식과 부모 Sharpe 구간별 차이",
        147,
    )
    mutate_row = origin[origin["group"] == "mutate"].iloc[0]
    crossover_row = origin[origin["group"] == "crossover"].iloc[0]
    add_body(
        doc,
        f"일반 변이로 만든 자식에서는 1축과 다축 차이가 {mutate_row['difference'] * 100:+.2f}퍼센트포인트로 거의 없었다. 교차 방식에서는 {crossover_row['difference'] * 100:+.2f}퍼센트포인트였다. 스윕에서 나온 다축은 109건뿐이었고 좋아진 사례가 하나도 없었다. 표본이 작아 안정적인 비교로 보기 어렵다. 전체 차이의 상당 부분이 자식을 만든 방식의 차이에서 생겼을 수 있다.",
    )
    add_body(
        doc,
        "부모 Sharpe가 가장 낮은 Q1에서 차이가 가장 컸다. Q2에서는 거의 사라졌고 Q3과 Q4에서는 다시 1축이 앞섰다. 부모 성적이 높아질수록 차이가 일정하게 커지거나 작아지는 모양은 아니다. 변이 폭을 정할 때 부모 성적만 보면 부족하다. 어떤 이유로 실패했는지와 최근 개선이 멈췄는지도 함께 봐야 한다.",
    )

    doc.add_heading("5.5 어느 축을 바꾸는지도 중요했다", level=2)
    gene_rows = []
    shown_genes = pd.concat([genes.head(6), genes.tail(4)]).drop_duplicates("gene")
    for row in shown_genes.itertuples():
        gene_rows.append(
            [
                row.gene,
                f"{int(row.n):,}",
                percent(row.improve_rate, 1),
                ci(row.ci_low, row.ci_high, 100, 1),
                number(row.delta_median, 2, sign=True),
            ]
        )
    add_caption(doc, "표 7. 1축 변이 안의 유전자별 탐색 결과")
    add_table(doc, ["유전자", "n", "개선률", "95% CI", "ΔSharpe 중앙값"], gene_rows)
    add_body(
        doc,
        "regime, trade_when, universe를 바꾼 사례는 개선률이 높았다. combine, hump, sign은 낮았다. 하지만 시스템이 축마다 서로 다른 상태의 부모를 골랐고 표본이 작은 축도 있다. 공정하게 무작위로 나눠 실험한 결과가 아니므로 이 순위만 보고 축별 정책을 바꾸면 안 된다. 다음 실험에서 먼저 확인할 축을 고르는 참고자료로 쓰는 편이 맞다.",
    )

    doc.add_heading("5.6 상관 비용은 뚜렷하게 늘지 않았다", level=2)
    add_body(
        doc,
        f"self-correlation 값이 남은 사례는 1축 {effects['self_corr']['single_n']:,}건, 다축 {effects['self_corr']['multi_n']:,}건뿐이었다. 중앙값은 각각 {effects['self_corr']['single']:.4f}, {effects['self_corr']['multi']:.4f}였다. 차이는 {effects['self_corr']['difference']:+.4f}(95% CI {ci(effects['self_corr']['ci_low'], effects['self_corr']['ci_high'], 1, 4)})였다. 이번에 확인할 수 있었던 사례에서는 1축이 기존 제출 알파와 더 비슷하다는 흔적이 나오지 않았다.",
    )
    add_body(
        doc,
        "self-correlation은 후보가 일정 단계까지 가야 기록된다. 값이 없는 사례가 우연히 빠졌다고 보기 어렵다. 따라서 상관 문제가 없다고 결론 내릴 수는 없다. 값이 남은 사례에서 1축과 다축의 차이를 찾지 못했다는 뜻이다.",
    )

    doc.add_heading("6. IOHprofiler 공개 문제 실험", level=1)
    doc.add_heading("6.1 최적점을 찾는 비율과 이동한 거리는 달랐다", level=2)
    external_rows = [
        ["정규화 진척도", number(external["progress_single"], 3), number(external["progress_multi"], 3), number(external["progress_difference"], 3, sign=True), ci(external["progress_ci_low"], external["progress_ci_high"], 1, 3)],
        ["최적점 도달률", percent(external["hit_single"], 1), percent(external["hit_multi"], 1), f"{external['hit_difference'] * 100:+.1f}pp", ci(external["hit_ci_low"], external["hit_ci_high"], 100, 1)],
        ["문제×차원 우세", f"{external['cells_single_better']}셀", f"{external['cells_multi_better']}셀", f"동률 {external['cells_tied']}셀", "총 50셀"],
    ]
    add_caption(doc, "표 8. 외부 벤치마크 종합 결과")
    add_table(doc, ["지표", "1축", "다축", "차이", "95% CI·범위"], external_rows)
    add_body(
        doc,
        f"1축은 최적점을 찾은 비율이 {percent(external['hit_single'], 1)}였다. 다축 {percent(external['hit_multi'], 1)}보다 {external['hit_difference'] * 100:.1f}퍼센트포인트 높다. 하지만 시작점에서 최적점 쪽으로 이동한 거리를 평균내면 1축은 {external['progress_single']:.3f}, 다축은 {external['progress_multi']:.3f}이었다. 1축과 다축의 차이에 대한 95% 신뢰구간은 {ci(external['progress_ci_low'], external['progress_ci_high'], 1, 3)}였다. 0이 들어 있어 평균 이동 거리에서 어느 쪽이 확실히 낫다고 말하기 어렵다.",
    )
    add_body(
        doc,
        "두 결과가 서로 모순되는 것은 아니다. 한 자리씩 고쳐도 되는 문제에서는 1축이 최적점까지 꾸준히 갔다. 여러 자리를 함께 바꿔야 다음 단계로 갈 수 있는 일부 문제에서는 다축이 훨씬 멀리 이동했다. 이런 소수 문제가 다축의 평균을 끌어올렸다.",
    )
    add_figure(
        doc,
        analysis_dir / "fig4_external_pbo_heatmap.png",
        "그림 4. IOH PBO 문제별 정규화 진척도 차이(1축−다축)",
        133,
    )

    doc.add_heading("6.2 문제 모양에 따라 유리한 방식이 달랐다", level=2)
    best = cells.sort_values("progress_difference", ascending=False).head(6)
    worst = cells.sort_values("progress_difference").head(6)
    cell_rows = []
    for direction, frame in (("1축 우세", best), ("다축 우세", worst)):
        for row in frame.itertuples():
            cell_rows.append(
                [
                    direction,
                    f"F{row.function_id} {row.function_name}",
                    str(row.dimension),
                    number(row.progress_single, 3),
                    number(row.progress_multi, 3),
                    number(row.progress_difference, 3, sign=True),
                ]
            )
    add_caption(doc, "표 9. 외부 문제에서 차이가 컸던 셀")
    add_table(doc, ["방향", "문제", "d", "1축", "다축", "차이"], cell_rows)
    add_body(
        doc,
        "OneMax와 크기 16의 LeadingOnes, Ising 계열에서는 1축이 강했다. OneMaxRuggedness2·3, OneMaxEpistasis, LeadingOnesEpistasis에서는 다축이 크게 앞섰다. 작은 변이가 언제나 좋은 것도 아니고 여러 축을 바꾸는 방식이 언제나 좋은 것도 아니다. 여러 자리를 함께 바꿔야 효과가 나는지, 중간에 넘기 어려운 벽이 있는지가 답을 바꾼다[5,6,7].",
    )

    doc.add_heading("7. 두 자료를 함께 읽기", level=1)
    synthesis_rows = [
        ["내부 단순 비교", "부모 대비 개선률", f"{effects['improved']['difference'] * 100:+.2f}pp", "1축 우세"],
        ["운영 조건을 맞춘 비교", "개선률·ΔSharpe", f"{regressions['adjusted_improved']['estimate'] * 100:+.2f}pp · {regressions['adjusted_delta_sharpe']['estimate']:+.3f}", "1축 우세"],
        ["같은 부모끼리 비교", "같은 부모", f"{regressions['sibling_fe_improved']['estimate'] * 100:+.2f}pp", "1축 우세"],
        ["내부 문턱", "S≥1.58 신규 도달", f"{regressions['adjusted_crossed_158']['estimate'] * 100:+.2f}pp", "우위 없음"],
        ["외부 도달", "최적점 도달률", f"{external['hit_difference'] * 100:+.1f}pp", "1축 우세"],
        ["외부 진척", "정규화 진척도", number(external['progress_difference'], 3, sign=True), "다축 쪽이나 불확실"],
    ]
    add_caption(doc, "표 10. 단일 가설에 대한 증거 통합")
    add_table(doc, ["분석", "평가 대상", "1축−다축", "판정"], synthesis_rows, widths=[35, 48, 40, 42])
    add_body(
        doc,
        "가설은 일부만 맞았다. GenomicWQB 기록에서는 1축이 부모보다 조금 나아질 가능성이 높았다. 운영 조건을 맞추거나 같은 부모의 자식끼리 비교해도 결과는 비슷했다. 하지만 2축도 1축만큼 좋았고 Sharpe 1.58을 새로 넘는 비율은 늘지 않았다. 공개 문제에서도 1축은 최적점에 자주 닿았다. 여러 축을 함께 바꿔야 풀리는 문제에서는 다축이 더 멀리 갔다.",
    )
    add_body(
        doc,
        "실무에서는 1축이 무조건 낫다고 정하면 안 된다. 성적이 좋은 부모를 조금 다듬을 때는 1~2축을 먼저 쓴다. 개선이 멈췄거나 여러 축을 함께 바꿔야 넘을 수 있는 벽이 보이면 3축 이상도 섞는다. 가까운 곳을 다듬는 일과 새로운 곳을 찾는 일을 같은 변이 폭 하나에 맡기지 않는 편이 좋다.",
    )

    doc.add_heading("8. 한계", level=1)
    add_body(
        doc,
        "첫째, 내부 배정은 무작위가 아니다. 스윕은 원래 1축을 많이 만들고 교차는 다축을 많이 만든다. 실패 유형에 따라 directed mutation이 고른 축도 달라진다. 회귀와 부모 고정효과는 관측된 구조를 조정하지만 숨은 배정 이유까지 없애지 못한다.",
    )
    add_body(
        doc,
        "둘째, 변경축 개수는 의미 거리와 다르다. decay 한 칸 이동과 fields 전체 교체를 각각 한 축으로 센다. 다축 안에서도 서로 보완하는 변화와 서로 상쇄하는 변화가 섞인다. 차기 원장은 축 수와 함께 축별 이동거리, 구조 변화량, AST 편집거리를 저장해야 한다.",
    )
    add_body(
        doc,
        "셋째, Sharpe는 IS 평가이고 반복 탐색은 데이터 스누핑 위험을 키운다[10]. 본 연구의 결과는 변이 연산자의 탐색 효율이지 실거래 수익의 증거가 아니다. S≥1.58도 공통 분석 문턱일 뿐 리전·지연·분류별 실제 제출 조건 전체를 대신하지 않는다.",
    )
    add_body(
        doc,
        "넷째, 외부 PBO는 이산 이진 문제다. 좌표 수와 상호작용을 통제하기에는 좋지만 알파 수식의 문법 제약, 연산자 의미, 데이터 상관을 재현하지 않는다. 외부 실험은 기제의 경계를 보여 줄 뿐 내부 효과 크기를 검증하지 않는다.",
    )
    add_body(
        doc,
        "다섯째, 보고서 개정 시점에 운영 시스템이 계속 변하고 있다. 스냅샷 이후의 후보와 정책은 분석에 들어오지 않는다. 재현할 때는 데이터 해시와 코드 버전을 함께 고정해야 한다.",
    )

    doc.add_heading("9. 운용 제안과 다음 검증", level=1)
    doc.add_heading("9.1 상황에 따라 변이 폭을 바꾸기", level=2)
    add_body(
        doc,
        "성적이 좋은 부모 주변을 다듬을 때는 1축과 2축을 기본으로 두는 편이 맞다. 최근에 자식 성적이 계속 좋아졌고 실패 원인이 한 기준에 모여 있다면 적게 바꿔 원인을 확인한다. 2축의 개선률도 1축보다 낮지 않았다. 함께 바꿀 때 도움이 되는 두 축까지 막을 이유는 없다.",
    )
    add_body(
        doc,
        "개선이 멈추면 다축을 탈출 수단으로 쓴다. 같은 부모에서 1~2축이 연속으로 실패하거나 여러 축을 함께 바꿔야 풀리는 문제로 보이면 3축 이상을 허용한다. 늘 같은 수를 바꾸기보다 현재 기록처럼 2~14축 사이에서 가끔 큰 값도 나오게 뽑을 수 있다. 큰 변이가 벽을 넘어야 하는 문제에서 유리하다는 연구와도 맞는다[6].",
    )
    doc.add_heading("9.2 다음에는 같은 부모로 세 방법을 공정하게 비교", level=2)
    add_body(
        doc,
        "이번 분석에는 운영 과정에서 생긴 선택 차이가 남아 있다. 다음 실험에서는 같은 부모가 1축, 2축, 3축 이상 자식을 모두 만들도록 실행 슬롯을 무작위로 나눠야 한다. 세 방법의 평가 횟수와 후보 수는 같게 둔다. 가장 중요한 지표는 부모보다 Sharpe가 좋아진 비율로 미리 정한다. Sharpe 1.58 신규 도달, self-correlation, 캐시를 뺀 실제 평가 시간도 시작 전에 함께 정해 둔다.",
    )
    trial_rows = [
        ["비교 묶음", "같은 parent_alpha_id와 같은 라운드"],
        ["처치군", "k=1 · k=2 · 경험적 k≥3"],
        ["가장 중요한 지표", "Sharpe_child > Sharpe_parent"],
        ["함께 볼 지표", "ΔSharpe · 문턱 도달 · self-correlation · 평가시간"],
        ["최소 표본", "독립 가정 4pp 검출 시 군당 약 1,176쌍"],
        ["오차 계산", "같은 부모를 한 묶음으로 계산"],
        ["중단 규칙", "오류율·캐시율 불균형 시 배정 로직 점검"],
    ]
    add_caption(doc, "표 11. 인과 검증을 위한 후속 실험안")
    add_table(doc, ["항목", "설계"], trial_rows, widths=[38, 125])
    add_body(
        doc,
        "기준 개선률을 12%로 놓고 4퍼센트포인트 차이를 찾으려면 방법마다 적어도 약 1,176쌍이 필요하다. 이 계산은 유의수준 0.05와 검정력 80%를 썼다. 실제 실험에서는 같은 부모에서 나온 자식끼리 닮는 정도와 중간에 빠지는 사례를 고려해 더 많이 모아야 한다. 이 실험이 끝나기 전에는 지금 나온 차이가 변이 폭 때문에 생겼다고 단정하지 않는다.",
    )

    doc.add_heading("10. 결론", level=1)
    add_body(
        doc,
        f"최신 GenomicWQB 기록 {snapshot['alphas']:,}건에서 실제로 새 평가를 돌린 부모·자식 {flow['primary_evaluated_lineage_pairs']:,}쌍을 골랐다. 1축은 다축보다 부모 Sharpe를 높일 가능성이 {effects['improved']['difference'] * 100:.2f}퍼센트포인트 컸다. 운영 조건을 맞추거나 같은 부모의 자식끼리 비교해도 약 4~5퍼센트포인트 차이가 났다. 좋은 부모 주변을 조금 다듬는 상황에서는 가설이 맞았다.",
    )
    add_body(
        doc,
        "하지만 1축이 언제나 낫지는 않았다. 2축은 1축과 비슷했고 Sharpe 기준을 새로 넘는 비율도 좋아지지 않았다. 공개 PBO 실험에서는 1축이 최적점을 더 자주 찾았다. 반대로 여러 자리를 함께 바꿔야 풀리는 문제에서는 다축이 더 멀리 갔다. 적게 바꾸는 방식은 좋은 후보를 다듬는 데 알맞고 많이 바꾸는 방식은 막힌 곳을 벗어나는 데 알맞다.",
    )
    add_body(
        doc,
        "운용 방법은 간단하다. 성적이 좋은 부모를 다듬을 때는 1~2축을 먼저 쓴다. 개선이 멈추거나 같은 실패가 반복되면 3축 이상을 섞는다. 다음 기록부터 같은 부모에 세 방법을 무작위로 나누고 축마다 얼마나 크게 바뀌었는지도 남기면 어느 방식이 원인인지 더 분명하게 확인할 수 있다.",
    )


def build_academic_report(doc: Document, analysis_dir: Path, summary: dict) -> None:
    snapshot = summary["snapshot"]
    flow = summary["cohort_flow"]
    effects = summary["effects"]
    regressions = summary["regressions"]
    external = summary["external"]
    by_k = pd.read_csv(analysis_dir / "internal_by_mutation_width.csv")
    hetero = pd.read_csv(analysis_dir / "internal_heterogeneity.csv")
    genes = pd.read_csv(analysis_dir / "internal_single_gene_results.csv")
    cells = pd.read_csv(analysis_dir / "external_benchmark_cells.csv")
    runs = pd.read_csv(analysis_dir / "external_benchmark_runs.csv")

    doc.add_heading("1. 서론", level=1)
    add_body(
        doc,
        "유전 알고리즘으로 알파를 찾을 때는 기존 후보의 어느 부분을 얼마나 바꿀지 정해야 한다. 알파 하나에는 데이터필드, 변환 연산자, 룩백, 중립화, 감쇠, 유니버스 같은 여러 설정이 들어간다. 한 번에 많이 바꾸면 새로운 영역으로 멀리 갈 수 있다. 한 축만 바꾸면 좋은 부모의 구조를 거의 그대로 둔 채 주변을 꼼꼼히 살필 수 있다. 평가에 시간과 실행 슬롯이 드는 WorldQuant BRAIN에서는 이 선택이 곧 연구 예산의 배분 문제가 된다.",
    )
    add_body(
        doc,
        "기존 보고서는 변이 폭과 다목적 선택을 함께 다뤘다. 이번 연구는 질문을 하나로 줄였다. 좋은 후보 주변을 다듬는 단계에서 1축 변이는 다축 변이보다 효과적인가. 답을 얻는 순서도 바꿨다. 실제 알파 기록부터 보는 대신 조건을 정확히 맞출 수 있는 공개 최적화 문제에서 작동 원리를 먼저 확인한다. 그 다음 GenomicWQB 운영 기록에서 같은 현상이 나타나는지 검정한다.",
    )
    add_body(
        doc,
        "이 순서는 두 자료의 역할을 분명히 한다. 공개 문제는 다른 조건을 같게 두고 바꾸는 축 수만 달리할 수 있어 원리를 확인하기 좋다. GenomicWQB는 실제 퀀트 알파 탐색에서 그 원리가 유효한지 보여 준다. 성격이 다른 두 자료가 같은 방향을 보이면 한 자료에만 기대는 결론보다 강하다. 다만 외부 실험이 알파 수익성을 대신하지 않으며 내부 기록도 무작위 실험은 아니라는 경계는 끝까지 유지한다.",
    )
    add_note(
        doc,
        "이 보고서에서 1축은 genes_changed에 기록된 유전자가 정확히 하나인 경우다. 다축은 둘 이상이다. 축 수는 바꾼 항목의 개수이며 변화량이나 수식 의미의 거리를 뜻하지 않는다.",
    )

    doc.add_heading("2. 이론적 배경과 연구가설", level=1)
    doc.add_heading("2.1 작은 변이는 좋은 후보 주변을 세밀하게 찾는다", level=2)
    add_body(
        doc,
        "고전적인 유전 알고리즘은 선택, 교차, 변이를 반복하며 후보를 개선한다(Holland, 1975; Koza, 1992). 변이 폭은 가까운 곳을 다듬는 활용과 새로운 곳을 찾는 탐색 사이의 균형을 정한다. OneMax처럼 각 자리를 따로 고쳐도 성적이 좋아지는 문제에서는 한 자리 변이가 강하다. 이미 좋은 후보에서 여러 자리를 함께 바꾸면 맞았던 자리까지 다시 틀릴 가능성이 커지기 때문이다(Buzdalov and Doerr, 2020).",
    )
    add_body(
        doc,
        "이 논리는 퀀트 알파에도 적용된다. 유망한 부모의 구조를 보존하면 어느 변경이 Sharpe를 움직였는지 추적하기 쉽고 실패 원인도 한 축 가까이 좁혀진다. 반대로 데이터필드와 결합 방식처럼 여러 요소가 함께 작동하면 한 축만 바꿔서는 벽을 넘기 어렵다. 큰 변이율이 점프형 문제에서 유리하다는 결과와 축 사이 상호작용이 알고리즘 성능을 바꾼다는 연구는 이 예외를 설명한다(Doerr et al., 2017; Salomon, 1996).",
    )

    doc.add_heading("2.2 연구가설은 국소 탐색에 한정한다", level=2)
    add_body(
        doc,
        "1축이 모든 문제에서 다축보다 낫다는 주장은 세우지 않는다. 좋은 부모 주변을 다듬는 국소 탐색에서 1축 변이가 다축보다 부모의 성적을 높일 가능성이 크다는 명제를 H1로 둔다. 차이가 없거나 다축이 더 낫다는 기준은 H0이다. 외부 실험에서는 최적점 도달률을 가장 중요한 지표로 삼고 이동 거리도 함께 본다. 내부 분석에서는 자식 Sharpe가 부모보다 높아졌는지를 가장 중요한 지표로 둔다.",
    )
    hypothesis_rows = [
        ["가설 H1", "국소 탐색에서 1축 변이가 다축보다 부모 성적을 높일 가능성이 크다"],
        ["외부의 핵심 지표", "같은 시작점과 평가예산에서 최적점에 도달했는지"],
        ["외부의 보조 지표", "시작점에서 최적점까지의 거리를 얼마나 줄였는지"],
        ["내부의 핵심 지표", "Sharpe_child > Sharpe_parent"],
        ["내부의 보조 지표", "ΔSharpe, S≥1.58 신규 도달, self-correlation"],
        ["판정 범위", "좋은 후보를 다듬는 활용 단계에 한정"],
    ]
    add_caption(doc, "<표 1> 단일 가설과 판정 기준")
    add_table(doc, ["항목", "내용"], hypothesis_rows, widths=[40, 120])

    doc.add_heading("3. IOHprofiler 공개 문제와 통제 실험", level=1)
    doc.add_heading("3.1 외부 자료", level=2)
    add_body(
        doc,
        "IOHprofiler의 Pseudo-Boolean Optimization(PBO) 25개 문제를 사용했다. PBO는 0과 1로 된 답을 조금씩 바꾸며 가장 좋은 값을 찾는 공개 문제 모음이다. 기본 함수인 OneMax와 LeadingOnes, 더미 변수와 중립성을 넣은 변형, 험준성과 축 간 상호작용을 넣은 변형, Ising과 N-Queens, NK Landscape 같은 조합 문제를 포함한다(Doerr et al., 2020; de Nobel et al., 2024). 실제 알파와 생김새는 다르지만 바꾸는 축 수를 정확히 맞출 수 있다.",
    )
    add_body(
        doc,
        f"각 문제는 크기 16과 100, 문제별 설정 1~3에서 실행했다. 조합마다 시작점 20개를 만들었다. 1축과 다축은 같은 시작점에서 출발하고 똑같이 20d번 평가받았다. 전체 비교는 {external['paired_runs']:,}쌍, 알고리즘 실행은 {external['paired_runs'] * 2:,}회, 함수평가는 6,960,000회다. 다축은 GenomicWQB에서 실제로 나타난 2~14축 비율에 맞춰 바꿀 자리 수를 뽑았다.",
    )

    doc.add_heading("3.2 실험 방법", level=2)
    add_body(
        doc,
        "두 알고리즘은 새 후보가 현재 후보보다 좋거나 같을 때만 다음 부모로 받아들이는 같은 규칙을 썼다. 다른 점은 한 번에 바꾸는 자리 수뿐이다. 1축은 매번 한 자리만 뒤집었다. 다축은 내부 기록의 변이 폭 분포에서 k를 뽑아 k자리를 뒤집었다. 하나의 시작점에서 나온 두 실행을 한 쌍으로 묶었기 때문에 시작점 운의 차이를 줄일 수 있다.",
    )
    external_design_rows = [
        ["문제", "IOHprofiler PBO 25종"],
        ["크기", "d=16, d=100"],
        ["문제별 설정", "instance 1, 2, 3"],
        ["반복", "조합마다 20회"],
        ["평가예산", "각 실행 20d"],
        ["1축", "매 평가마다 한 자리 변경"],
        ["다축", "GenomicWQB의 2~14축 비율에 따라 변경"],
        ["비교 단위", f"같은 문제·설정·시작점의 대응 실행 {external['paired_runs']:,}쌍"],
    ]
    add_caption(doc, "<표 2> 외부 통제 실험 설계")
    add_table(doc, ["항목", "설계"], external_design_rows, widths=[40, 120])

    doc.add_heading("4. 외부 실험 결과", level=1)
    doc.add_heading("4.1 1축은 최적점을 더 자주 찾았다", level=2)
    external_rows = [
        ["최적점 도달률", percent(external["hit_single"], 1), percent(external["hit_multi"], 1), f"{external['hit_difference'] * 100:+.1f}pp", ci(external["hit_ci_low"], external["hit_ci_high"], 100, 1)],
        ["평균 이동 거리", number(external["progress_single"], 3), number(external["progress_multi"], 3), number(external["progress_difference"], 3, sign=True), ci(external["progress_ci_low"], external["progress_ci_high"], 1, 3)],
        ["문제×크기 우세", f"{external['cells_single_better']}셀", f"{external['cells_multi_better']}셀", f"동률 {external['cells_tied']}셀", "총 50셀"],
    ]
    add_caption(doc, "<표 3> 외부 실험의 전체 결과")
    add_table(doc, ["지표", "1축", "다축", "차이", "95% CI·범위"], external_rows)
    add_body(
        doc,
        f"1축은 최적점을 찾은 비율이 {percent(external['hit_single'], 1)}로 다축 {percent(external['hit_multi'], 1)}보다 {external['hit_difference'] * 100:.1f}퍼센트포인트 높았다. 95% 신뢰구간도 {ci(external['hit_ci_low'], external['hit_ci_high'], 100, 1)}로 0보다 컸다. 한 자리씩 고쳐도 되는 문제에서는 1축이 좋은 답까지 안정적으로 모였다는 뜻이다.",
    )
    add_body(
        doc,
        f"평균 이동 거리는 1축 {external['progress_single']:.3f}, 다축 {external['progress_multi']:.3f}이었다. 1축−다축 차이는 {external['progress_difference']:+.3f}이지만 95% 신뢰구간 {ci(external['progress_ci_low'], external['progress_ci_high'], 1, 3)}에 0이 들어 있다. 전체 평균만으로 다축이 더 낫다고 단정하기 어렵다. 최적점 도달률과 이동 거리가 다른 답을 낸 이유는 문제 종류를 나누면 분명해진다.",
    )

    doc.add_heading("4.2 문제의 구조가 결과를 바꿨다", level=2)
    group_map = {
        **{item: "기본 문제" for item in (1, 2, 3)},
        **{item: "더미·중립성 변형" for item in (4, 5, 6, 11, 12, 13)},
        **{item: "험준·상호작용 변형" for item in (7, 8, 9, 10, 14, 15, 16, 17)},
        **{item: "조합 문제" for item in range(18, 26)},
    }
    runs["problem_group"] = runs["function_id"].map(group_map)
    paired = runs.pivot_table(
        index=["function_id", "dimension", "instance", "rep", "pair_seed", "problem_group"],
        columns="arm",
        values=["normalized_progress", "hit_optimum"],
    ).reset_index()
    group_rows = []
    for group in ("기본 문제", "더미·중립성 변형", "험준·상호작용 변형", "조합 문제"):
        frame = paired[paired["problem_group"] == group]
        progress_difference = (
            frame[("normalized_progress", "single")]
            - frame[("normalized_progress", "empirical_multi")]
        ).mean()
        hit_difference = (
            frame[("hit_optimum", "single")]
            - frame[("hit_optimum", "empirical_multi")]
        ).mean()
        group_rows.append(
            [group, f"{len(frame):,}", f"{hit_difference * 100:+.1f}pp", f"{progress_difference:+.3f}"]
        )
    add_caption(doc, "<표 4> 문제 종류별 1축−다축 차이")
    add_table(doc, ["문제 종류", "대응쌍", "최적점 도달률 차이", "평균 이동 거리 차이"], group_rows)
    add_body(
        doc,
        "기본 문제에서 1축의 최적점 도달률은 다축보다 69.2퍼센트포인트 높았고 이동 거리도 0.042만큼 앞섰다. 더미 변수와 중립성을 더한 문제에서는 도달률이 6.0퍼센트포인트 높았으나 이동 거리는 거의 같았다. 조합 문제에서도 1축이 두 지표 모두 앞섰다. 반면 험준성과 축 간 상호작용을 넣은 문제에서는 다축이 이동 거리에서 0.230만큼 앞섰다. 이 구분은 함수 이름을 보고 나눈 참고 분석이며 가설의 핵심 검정은 전체 대응쌍 결과다.",
    )
    add_figure(
        doc,
        analysis_dir / "fig4_external_pbo_heatmap.png",
        "<그림 1> IOH PBO 문제별 평균 이동 거리 차이(1축−다축)",
        108,
    )
    add_body(
        doc,
        "OneMax와 크기 16의 LeadingOnes, Ising 계열에서는 1축이 강했다. OneMaxRuggedness2·3과 OneMaxEpistasis, LeadingOnesEpistasis에서는 다축이 크게 앞섰다. 외부 실험은 1축이 좋은 후보 주변을 다듬고 최적점에 모이는 데 유리하다는 이론적 예측을 지지한다. 동시에 여러 축을 함께 바꿔야 넘을 수 있는 벽에서는 다축이 필요하다는 경계도 확인했다.",
    )

    doc.add_heading("5. GenomicWQB 자료와 분석 방법", level=1)
    doc.add_heading("5.1 운영 자료", level=2)
    add_body(
        doc,
        f"외부 실험에서 확인한 원리가 퀀트 알파에도 나타나는지 GenomicWQB 기록으로 검정했다. 자료는 {snapshot['max_kst'].replace('T', ' ')} 시점의 SQLite 복사본이며 기간은 {snapshot['min_kst'][:10]}부터 {snapshot['max_kst'][:10]}까지다. 알파 {snapshot['alphas']:,}건과 라운드 {snapshot['rounds']:,}건이 들어 있다. 서로 다른 표현식은 {snapshot['unique_codes']:,}종, 서로 다른 설정은 {snapshot['unique_settings']:,}종이다.",
    )
    add_body(
        doc,
        "부모와 자식의 Sharpe가 모두 있는 기록을 연결하고 genes_changed에서 바꾼 축 수를 셌다. 아무것도 바꾸지 않은 기록과 캐시를 다시 쓴 사례, 표현식과 설정이 모두 같은 사례는 제외했다. 최종 표본은 실제로 새 평가를 돌린 부모·자식 쌍이다.",
    )
    flow_rows = [
        ["부모·자식 Sharpe 연결", f"{flow['joined_parent_child_with_sharpe']:,}", "시작 표본"],
        ["변경축 0 제외", f"−{flow['excluded_zero_changed']:,}", "무변경·기록 불일치"],
        ["캐시 재사용 제외", f"−{flow['excluded_cached']:,}", "새 평가 없음"],
        ["완전 동일 설정 제외", f"−{flow['excluded_exact_config']:,}", "표현식과 설정 모두 동일"],
        ["시간순서 오류 제외", f"−{flow['excluded_invalid_time_order']:,}", "자식이 부모보다 이른 경우"],
        ["최종 분석 표본", f"{flow['primary_evaluated_lineage_pairs']:,}", f"1축 {flow['primary_single']:,} · 다축 {flow['primary_multi']:,}"],
    ]
    add_caption(doc, "<표 5> GenomicWQB 분석 표본을 고른 과정")
    add_table(doc, ["단계", "건수", "판정"], flow_rows, widths=[52, 28, 80])

    doc.add_heading("5.2 비교 방법", level=2)
    add_body(
        doc,
        "가장 중요한 지표는 자식 Sharpe가 부모보다 높아졌는지다. Sharpe 변화량과 1.58 신규 도달, self-correlation도 함께 봤다. 같은 부모에서 나온 자식은 출발점과 실패 원인이 비슷하다. 모든 자식을 서로 무관한 사례로 보면 오차 범위가 실제보다 좁아질 수 있어 부모를 한 묶음으로 삼아 2,000번 다시 뽑았다. 이 방법을 부모 단위 군집 부트스트랩이라고 한다(Cameron and Miller, 2015).",
    )
    add_body(
        doc,
        "운영 조건의 차이도 통계적으로 맞췄다. 부모 Sharpe와 그 제곱, 생성 세대와 방식, 리전, 유니버스, 지연, 평가 주차를 포함했다. 1축과 다축 자식을 모두 만든 부모만 골라 자식끼리 비교하는 계산도 했다. 1축과 다축 자식을 모두 가진 부모는 "
        f"{regressions['mixed_parent_count']:,}개였고 자식은 {regressions['mixed_parent_rows']:,}건이었다. 이 비교는 부모 자체의 품질 차이를 줄이지만 시스템이 어떤 변이 폭을 골랐는지에 관한 숨은 이유까지 없애지는 못한다.",
    )

    doc.add_heading("6. GenomicWQB 실증 결과", level=1)
    doc.add_heading("6.1 1~2축과 3축 이상 사이에서 차이가 벌어졌다", level=2)
    by_k_rows = []
    for row in by_k.itertuples():
        by_k_rows.append(
            [
                f"{row.k_group}축",
                f"{int(row.n):,}",
                percent(row.improve_rate, 2),
                ci(row.improve_ci_low, row.improve_ci_high, 100, 2),
                number(row.delta_median, 2, sign=True),
                percent(row.crossed_158_rate, 2),
            ]
        )
    add_caption(doc, "<표 6> 바꾼 축 수별 GenomicWQB 결과")
    add_table(doc, ["변이 폭", "n", "개선률", "95% CI", "ΔSharpe 중앙값", "S≥1.58 신규 도달"], by_k_rows)
    add_figure(
        doc,
        analysis_dir / "fig1_internal_dose_response.png",
        "<그림 2> 바꾼 축 수와 부모 대비 개선률·Sharpe 변화량",
        150,
    )
    add_body(
        doc,
        f"1축 개선률은 {percent(by_k.iloc[0]['improve_rate'], 2)}, 2축은 {percent(by_k.iloc[1]['improve_rate'], 2)}였다. 두 구간의 오차 범위는 크게 겹쳤다. 3축은 {percent(by_k.iloc[2]['improve_rate'], 2)}, 4축은 {percent(by_k.iloc[3]['improve_rate'], 2)}로 내려갔고 5축과 6축 이상은 약 6%였다. 1축만 특별히 좋다기보다 1~2축처럼 적게 바꿀 때와 3축 이상을 넓게 바꿀 때의 차이가 뚜렷하다.",
    )
    add_body(
        doc,
        "모든 구간에서 Sharpe 변화량의 중앙값은 0보다 작았다. 새 자식은 대체로 부모보다 나빴다. 1축이 평균 성적을 높였다는 뜻은 아니다. 다축보다 실패가 적고 떨어지는 폭도 작았다는 뜻이다.",
    )

    doc.add_heading("6.2 세 가지 비교에서 비슷한 차이가 나왔다", level=2)
    effect_rows = [
        ["단순 비교", f"{effects['improved']['difference'] * 100:+.2f}pp", ci(effects['improved']['ci_low'], effects['improved']['ci_high'], 100, 2), f"n={effects['improved']['single_n'] + effects['improved']['multi_n']:,}"],
        ["운영 조건을 맞춘 비교", f"{regressions['adjusted_improved']['estimate'] * 100:+.2f}pp", ci(regressions['adjusted_improved']['ci_low'], regressions['adjusted_improved']['ci_high'], 100, 2), f"p={regressions['adjusted_improved']['p_value']:.4g}"],
        ["같은 부모의 자식 비교", f"{regressions['sibling_fe_improved']['estimate'] * 100:+.2f}pp", ci(regressions['sibling_fe_improved']['ci_low'], regressions['sibling_fe_improved']['ci_high'], 100, 2), f"부모 {regressions['mixed_parent_count']:,}개"],
        ["부모마다 같은 무게", f"{regressions['sibling_equal_parent_improved']['estimate'] * 100:+.2f}pp", ci(regressions['sibling_equal_parent_improved']['ci_low'], regressions['sibling_equal_parent_improved']['ci_high'], 100, 2), "부모당 같은 가중치"],
    ]
    add_caption(doc, "<표 7> 1축 변이의 부모 대비 개선률 차이")
    add_table(doc, ["비교 방법", "1축−다축", "95% CI", "표본·검정"], effect_rows)
    add_figure(
        doc,
        analysis_dir / "fig2_internal_effect_estimates.png",
        "<그림 3> 비교 방법별 1축 변이의 개선률 차이",
        132,
    )
    add_body(
        doc,
        f"그대로 비교하면 1축 개선률이 {effects['improved']['difference'] * 100:.2f}퍼센트포인트 높았다. 운영 조건을 맞추면 {regressions['adjusted_improved']['estimate'] * 100:.2f}퍼센트포인트, 같은 부모의 자식끼리 비교하면 {regressions['sibling_fe_improved']['estimate'] * 100:.2f}퍼센트포인트 높았다. 세 계산의 방향과 크기가 비슷하다. 기록으로 확인할 수 있는 부모 품질과 평가 시점만으로는 이 차이를 설명하기 어렵다.",
    )
    add_body(
        doc,
        f"Sharpe 변화량도 같은 방향이었다. 단순 평균 차이는 {effects['delta_sharpe']['difference']:+.3f}, 운영 조건을 맞춘 차이는 {regressions['adjusted_delta_sharpe']['estimate']:+.3f}(95% CI {ci(regressions['adjusted_delta_sharpe']['ci_low'], regressions['adjusted_delta_sharpe']['ci_high'], 1, 3)}), 같은 부모의 자식끼리 비교한 차이는 {regressions['sibling_fe_delta_sharpe']['estimate']:+.3f}이었다. 1축 자식도 평균적으로 부모보다 나빴지만 다축보다 덜 떨어졌다.",
    )

    doc.add_heading("6.3 최종 기준 도달은 좋아지지 않았다", level=2)
    add_body(
        doc,
        f"Sharpe 1.58을 새로 넘은 비율은 1축 {percent(effects['crossed_158']['single'], 2)}, 다축 {percent(effects['crossed_158']['multi'], 2)}였다. 단순 차이는 {effects['crossed_158']['difference'] * 100:+.2f}퍼센트포인트였고 95% 신뢰구간은 {ci(effects['crossed_158']['ci_low'], effects['crossed_158']['ci_high'], 100, 2)}였다. 운영 조건을 맞춘 차이도 {regressions['adjusted_crossed_158']['estimate'] * 100:+.2f}퍼센트포인트에 그쳤다. 1축은 작은 개선을 자주 만들지만 낮은 부모를 한 번에 기준 밖으로 끌어올리는 힘은 약하다.",
    )
    add_body(
        doc,
        f"self-correlation이 남은 사례는 1축 {effects['self_corr']['single_n']:,}건, 다축 {effects['self_corr']['multi_n']:,}건이었다. 중앙값은 각각 {effects['self_corr']['single']:.4f}, {effects['self_corr']['multi']:.4f}였고 차이는 {effects['self_corr']['difference']:+.4f}(95% CI {ci(effects['self_corr']['ci_low'], effects['self_corr']['ci_high'], 1, 4)})였다. 값이 남은 사례에서는 1축이 기존 제출 알파와 더 비슷하다는 흔적이 나오지 않았다. 다만 이 값은 후보가 일정 단계까지 가야 기록되므로 상관 문제가 없다고 결론 내릴 수는 없다.",
    )

    doc.add_heading("6.4 자식을 만든 방식과 부모 성적에 따라 달랐다", level=2)
    origin = hetero[hetero["dimension"] == "origin"].copy()
    quartile = hetero[hetero["dimension"] == "parent_quartile"].copy()
    hetero_rows = []
    for row in pd.concat([origin, quartile]).itertuples():
        label = {"crossover": "교차", "mutate": "일반 변이", "sweep": "스윕"}.get(row.group, row.group)
        hetero_rows.append(
            ["생성 방식" if row.dimension == "origin" else "부모 Sharpe", label, f"{int(row.n):,}", percent(row.single_rate, 1), percent(row.multi_rate, 1), f"{row.difference * 100:+.1f}pp", ci(row.ci_low, row.ci_high, 100, 1)]
        )
    add_caption(doc, "<표 8> 생성 방식과 부모 Sharpe 구간별 개선률 차이")
    add_table(doc, ["구분", "집단", "n", "1축", "다축", "차이", "95% CI"], hetero_rows)
    add_figure(
        doc,
        analysis_dir / "fig3_internal_heterogeneity.png",
        "<그림 4> 생성 방식과 부모 Sharpe 구간별 차이",
        142,
    )
    mutate_row = origin[origin["group"] == "mutate"].iloc[0]
    crossover_row = origin[origin["group"] == "crossover"].iloc[0]
    add_body(
        doc,
        f"일반 변이에서 1축과 다축 차이는 {mutate_row['difference'] * 100:+.2f}퍼센트포인트로 거의 없었다. 교차 방식에서는 {crossover_row['difference'] * 100:+.2f}퍼센트포인트였다. 스윕에서 나온 다축은 109건뿐이었고 좋아진 사례가 하나도 없어 비교가 불안정하다. 전체 차이의 일부가 자식을 만든 방식에서 생겼을 가능성이 남는다.",
    )

    doc.add_heading("6.5 어느 축을 바꾸는지도 중요했다", level=2)
    shown_genes = pd.concat([genes.head(6), genes.tail(4)]).drop_duplicates("gene")
    gene_rows = []
    for row in shown_genes.itertuples():
        gene_rows.append([row.gene, f"{int(row.n):,}", percent(row.improve_rate, 1), ci(row.ci_low, row.ci_high, 100, 1), number(row.delta_median, 2, sign=True)])
    add_caption(doc, "<표 9> 1축 변이 안의 유전자별 참고 결과")
    add_table(doc, ["유전자", "n", "개선률", "95% CI", "ΔSharpe 중앙값"], gene_rows)
    add_body(
        doc,
        "regime, trade_when, universe를 바꾼 사례는 개선률이 높았고 combine, hump, sign은 낮았다. 시스템이 축마다 서로 다른 부모를 골랐고 표본이 작은 축도 있다. 무작위 실험이 아니므로 이 순위만 보고 축별 정책을 바꾸면 안 된다. 후속 실험에서 먼저 확인할 축을 고르는 참고자료로 사용한다.",
    )

    doc.add_heading("7. 강건성 검토와 증거 통합", level=1)
    synthesis_rows = [
        ["외부 전체", "최적점 도달률", f"{external['hit_difference'] * 100:+.1f}pp", "1축 우세"],
        ["외부 전체", "평균 이동 거리", number(external['progress_difference'], 3, sign=True), "차이 불확실"],
        ["WQB 단순 비교", "부모 대비 개선률", f"{effects['improved']['difference'] * 100:+.2f}pp", "1축 우세"],
        ["WQB 조건 보정", "개선률·ΔSharpe", f"{regressions['adjusted_improved']['estimate'] * 100:+.2f}pp · {regressions['adjusted_delta_sharpe']['estimate']:+.3f}", "1축 우세"],
        ["WQB 같은 부모", "같은 부모의 자식", f"{regressions['sibling_fe_improved']['estimate'] * 100:+.2f}pp", "1축 우세"],
        ["WQB 최종 기준", "S≥1.58 신규 도달", f"{regressions['adjusted_crossed_158']['estimate'] * 100:+.2f}pp", "우위 없음"],
    ]
    add_caption(doc, "<표 10> 외부 실험과 GenomicWQB 결과의 통합")
    add_table(doc, ["자료·방법", "평가 대상", "1축−다축", "판정"], synthesis_rows)
    add_body(
        doc,
        "공개 문제와 실제 알파 기록은 서로 다른 약점을 가진다. 공개 문제는 조건을 정확히 맞출 수 있지만 알파의 문법과 경제적 의미를 담지 못한다. WQB 기록은 실제 알파 탐색을 보여 주지만 변이 폭이 무작위로 정해지지 않았다. 그럼에도 공개 문제의 최적점 도달률과 WQB의 부모 대비 개선률이 같은 방향을 보였다. 1축이 좋은 후보 주변을 다듬는 데 유리하다는 결론은 한 종류의 자료에만 기대지 않는다.",
    )
    add_body(
        doc,
        "결론의 범위는 분명하다. 1축이 모든 탐색 단계에서 가장 좋다는 뜻은 아니다. 2축은 WQB에서 1축과 비슷했고 험준성과 축 간 상호작용이 큰 공개 문제에서는 다축이 더 멀리 갔다. 강건한 결론은 유망한 후보를 조금씩 개선하는 국소 탐색에서 1축 또는 2축의 좁은 변이가 유효하다는 데 한정된다. 이 원리가 통제된 공개 문제와 퀀트 알파의 실제 운영 기록에서 함께 나타났다.",
    )

    doc.add_heading("8. 운용 제안과 후속 검증", level=1)
    doc.add_heading("8.1 탐색 상태에 따라 변이 폭을 바꾼다", level=2)
    add_body(
        doc,
        "성적이 좋은 부모 주변을 다듬을 때는 1축과 2축을 기본으로 둔다. 최근 자식 성적이 계속 좋아졌고 실패 원인이 한 기준에 모여 있다면 적게 바꿔 원인을 확인한다. 2축 개선률도 1축보다 낮지 않았으므로 함께 바꿀 때 도움이 되는 두 축까지 막을 이유는 없다.",
    )
    add_body(
        doc,
        "개선이 멈추면 다축을 탈출 수단으로 쓴다. 같은 부모에서 1~2축이 연속으로 실패하거나 여러 축을 함께 바꿔야 풀리는 문제로 보이면 3축 이상을 허용한다. 늘 같은 수를 바꾸기보다 현재 기록처럼 2~14축 사이에서 가끔 큰 값도 나오게 뽑을 수 있다.",
    )

    doc.add_heading("8.2 같은 부모로 세 방법을 무작위 비교한다", level=2)
    add_body(
        doc,
        "다음 실험에서는 같은 부모가 1축, 2축, 3축 이상 자식을 모두 만들도록 실행 슬롯을 무작위로 나눈다. 세 방법의 평가 횟수와 후보 수는 같게 둔다. 가장 중요한 지표는 부모보다 Sharpe가 좋아진 비율로 미리 정한다. Sharpe 1.58 신규 도달, self-correlation, 캐시를 뺀 실제 평가 시간도 시작 전에 정해 둔다.",
    )
    trial_rows = [
        ["비교 묶음", "같은 parent_alpha_id와 같은 라운드"],
        ["세 방법", "k=1 · k=2 · 경험적 k≥3"],
        ["가장 중요한 지표", "Sharpe_child > Sharpe_parent"],
        ["함께 볼 지표", "ΔSharpe · 기준 도달 · self-correlation · 평가시간"],
        ["최소 표본", "독립 가정 4pp 차이 검출 시 방법마다 약 1,176쌍"],
        ["오차 계산", "같은 부모를 한 묶음으로 계산"],
        ["중단 기준", "오류율·캐시율 불균형이 생기면 배정 로직 점검"],
    ]
    add_caption(doc, "<표 11> 인과 검증을 위한 후속 실험")
    add_table(doc, ["항목", "설계"], trial_rows, widths=[42, 118])
    add_body(
        doc,
        "기준 개선률을 12%로 놓고 4퍼센트포인트 차이를 찾으려면 방법마다 적어도 약 1,176쌍이 필요하다. 유의수준 0.05와 검정력 80%를 쓴 계산이다. 실제 실험에서는 같은 부모에서 나온 자식끼리 닮는 정도와 중간에 빠지는 사례를 고려해 더 많이 모아야 한다.",
    )

    doc.add_heading("9. 연구의 한계", level=1)
    add_body(
        doc,
        "첫째, WQB에서 1축과 다축은 무작위로 배정되지 않았다. 스윕은 1축을 많이 만들고 교차는 다축을 많이 만든다. 통계 계산으로 기록에 남은 차이는 줄였지만 시스템이 어떤 변이 폭을 골랐는지에 관한 숨은 이유까지 없앨 수는 없다.",
    )
    add_body(
        doc,
        "둘째, 바꾼 축 수가 같다고 변화의 크기까지 같은 것은 아니다. decay를 한 칸 옮기는 일과 fields 전체를 바꾸는 일을 모두 1축으로 센다. 다음 기록에는 축 수와 함께 축별 이동 거리, 수식 구조 변화량, AST 편집거리를 남길 필요가 있다.",
    )
    add_body(
        doc,
        "셋째, 여기서 쓴 Sharpe는 IS 평가값이다. 같은 자료로 여러 번 탐색하면 우연히 좋아 보이는 알파를 고를 위험이 커진다(Bailey et al., 2014). 이번 결과는 탐색 방식의 효율을 보여 줄 뿐 실거래 수익을 보장하지 않는다. Sharpe 1.58도 분석을 위한 공통 기준이며 실제 제출 조건 전체를 대신하지 않는다.",
    )
    add_body(
        doc,
        "넷째, 외부 PBO는 0과 1로 만든 최적화 문제다. 바꾸는 자리 수와 자리 사이의 영향을 정확히 맞추기에는 좋지만 알파 수식의 문법과 연산자 의미, 데이터 상관관계는 재현하지 못한다. 외부 실험은 작동 원리의 범위를 보여 줄 뿐 WQB에서 나온 효과 크기를 대신 검증하지 않는다.",
    )
    add_body(
        doc,
        "다섯째, 운영 시스템은 계속 바뀐다. 자료를 복사한 시점 뒤에 만든 후보와 새 정책은 분석에 들어 있지 않다. 같은 분석을 다시 하려면 데이터 해시와 코드 버전을 함께 고정해야 한다.",
    )

    doc.add_heading("10. 결론", level=1)
    add_body(
        doc,
        f"공개 PBO 문제 {external['paired_runs']:,}쌍을 먼저 비교한 결과 1축은 다축보다 최적점을 {external['hit_difference'] * 100:.1f}퍼센트포인트 더 자주 찾았다. 기본 문제와 한 자리씩 고쳐도 되는 문제에서 차이가 컸다. 험준성과 축 간 상호작용이 큰 문제에서는 다축이 더 멀리 갔다. 외부 실험은 1축을 모든 상황의 우월한 방식이 아니라 좋은 후보 주변을 다듬는 방식으로 설명한다.",
    )
    add_body(
        doc,
        f"이 작동 원리는 퀀트 알파에서도 나타났다. GenomicWQB 기록 {snapshot['alphas']:,}건에서 실제 부모·자식 {flow['primary_evaluated_lineage_pairs']:,}쌍을 분석했을 때 1축은 다축보다 부모 Sharpe를 높일 가능성이 {effects['improved']['difference'] * 100:.2f}퍼센트포인트 컸다. 운영 조건을 맞추거나 같은 부모의 자식끼리 비교해도 약 4~5퍼센트포인트 차이가 유지됐다. 통제된 공개 문제와 실제 알파 기록이 같은 방향을 보였으므로 1축 변이는 퀀트 알파의 국소 탐색에서도 유효하다는 결론을 얻는다.",
    )
    add_body(
        doc,
        "다만 정확히 한 축만 고집할 근거는 없다. WQB에서 2축은 1축과 비슷했고 최종 Sharpe 기준 도달률은 좋아지지 않았다. 운용할 때는 좋은 부모를 다듬는 단계에서 1~2축을 먼저 쓰고 개선이 멈추거나 구조적 실패가 반복되면 3축 이상을 섞는다. 다음 단계는 같은 부모에서 세 변이 폭을 무작위로 비교해 지금의 강건한 관측 결론을 인과 결론으로 옮기는 일이다.",
    )


def add_references(doc: Document) -> None:
    doc.add_heading("11. 참고문헌", level=1)
    references = [
        ("[1] Holland, J. H. (1975). Adaptation in Natural and Artificial Systems. University of Michigan Press.", None),
        ("[2] Koza, J. R. (1992). Genetic Programming: On the Programming of Computers by Means of Natural Selection. MIT Press.", None),
        ("[3] Doerr, C., Ye, F., Horesh, N., Wang, H., Shir, O. M., & Bäck, T. (2020). Benchmarking discrete optimization heuristics with IOHprofiler. Applied Soft Computing, 88, 106027.", "https://doi.org/10.1016/j.asoc.2019.106027"),
        ("[4] de Nobel, J., Ye, F., Vermetten, D., Wang, H., Doerr, C., & Bäck, T. (2024). IOHexperimenter: Benchmarking Platform for Iterative Optimization Heuristics. Evolutionary Computation, 32(3), 205–210.", "https://doi.org/10.1162/evco_a_00342"),
        ("[5] Buzdalov, M., & Doerr, C. (2020). Optimal Mutation Rates for the (1+λ) EA on OneMax. In Parallel Problem Solving from Nature – PPSN XVI, 574–587.", "https://doi.org/10.1007/978-3-030-58115-2_40"),
        ("[6] Doerr, B., Le, H. P., Makhmara, R., & Nguyen, T. D. (2017). Fast Genetic Algorithms. GECCO ’17, 777–784.", "https://doi.org/10.1145/3071178.3071301"),
        ("[7] Salomon, R. (1996). Re-evaluating genetic algorithm performance under coordinate rotation of benchmark functions. BioSystems, 39(3), 263–278.", "https://doi.org/10.1016/0303-2647(96)01621-8"),
        ("[8] Cameron, A. C., & Miller, D. L. (2015). A Practitioner’s Guide to Cluster-Robust Inference. Journal of Human Resources, 50(2), 317–372.", "https://doi.org/10.3368/jhr.50.2.317"),
        ("[9] Bergstra, J., & Bengio, Y. (2012). Random Search for Hyper-Parameter Optimization. Journal of Machine Learning Research, 13, 281–305.", "https://www.jmlr.org/papers/v13/bergstra12a.html"),
        ("[10] Bailey, D. H., Borwein, J. M., López de Prado, M., & Zhu, Q. J. (2014). Pseudo-Mathematics and Financial Charlatanism: The Effects of Backtest Overfitting on Out-of-Sample Performance. Notices of the AMS, 61(5), 458–471.", "https://doi.org/10.1090/noti1105"),
        ("[11] IOHprofiler. IOHdata: official benchmark data sets generated from IOHexperimenter. GitHub commit a3ecaac6c06aa97190249c0c2c74b6a3aac14304, 2026-08-21 인출.", "https://github.com/IOHprofiler/IOHdata"),
        ("[12] IOHprofiler. IOHexperimenter Python Interface and PBO Problem Suite. ioh 0.3.22, 2026-08-21 인출.", "https://iohprofiler.github.io/IOHexperimenter/python.html"),
        ("[13] WorldQuant BRAIN. Consultant Dos and Don’ts; Simulation Settings; Data and Operator Documentation. GenomicWQB 로컬 수집본, 2026-08-21 확인.", None),
        ("[14] GenomicWQB 운영 데이터베이스. alphas·rounds·submit_attempts, 2026-08-21 16:09:09 KST 스냅샷.", None),
    ]
    references.sort(key=lambda item: re.sub(r"^\[\d+\]\s*", "", item[0]).casefold())
    for text, url in references:
        text = re.sub(r"^\[\d+\]\s*", "", text)
        paragraph = doc.add_paragraph(style="Reference")
        paragraph.add_run(text)
        if url:
            paragraph.add_run(" ")
            add_hyperlink(paragraph, url, url)


def add_appendix(doc: Document, analysis_dir: Path, summary: dict) -> None:
    by_k = pd.read_csv(analysis_dir / "internal_by_mutation_width.csv")
    cells = pd.read_csv(analysis_dir / "external_benchmark_cells.csv")
    doc.add_heading("부록 A. 변수와 집계식", level=1)
    variable_rows = [
        ["k", "len(json.loads(genes_changed))", "동시에 바뀐 선언 유전자 수"],
        ["single", "1(k=1)", "1축 처치 지시자"],
        ["improved", "1(Sharpe_child>Sharpe_parent)", "주 결과"],
        ["delta_sharpe", "Sharpe_child−Sharpe_parent", "연속 보조 결과"],
        ["crossed_158", "1(parent<1.58≤child)", "공통 분석 문턱 신규 도달"],
        ["self_corr", "alphas.self_corr 또는 metrics", "기존 제출 풀과 최대 상관"],
        ["exact_config", "code_hash와 settings_fp 동시 일치", "제외 규칙"],
        ["cached", "alphas.cached=1", "신규 실평가가 아닌 행"],
    ]
    add_table(doc, ["변수", "정의", "용도"], variable_rows, widths=[35, 65, 65])

    doc.add_heading("부록 B. 다축 변이 강도 분포", level=1)
    distribution = summary["multi_strength_distribution"]
    dist_rows = [[f"{key}축", percent(float(value), 2)] for key, value in distribution.items()]
    add_body(
        doc,
        "외부 벤치마크의 다축 알고리즘은 아래 내부 분포에서 k를 복원추출했다. 차원이 k보다 작으면 차원 수로 제한했다. 이번 실험의 차원은 16과 100이라 실제 제한은 발생하지 않았다.",
    )
    add_table(doc, ["k", "내부 다축 표본 비중"], dist_rows, widths=[35, 55])

    doc.add_heading("부록 C. IOH PBO 문제 목록", level=1)
    problem_rows = []
    for function_id, frame in cells.groupby("function_id"):
        row = frame.iloc[0]
        dims = " · ".join(
            f"d={int(item.dimension)} {item.progress_difference:+.3f}"
            for item in frame.sort_values("dimension").itertuples()
        )
        problem_rows.append([f"F{int(function_id)}", row["function_name"], dims])
    add_table(doc, ["ID", "문제", "정규화 진척도 차이(1축−다축)"], problem_rows, widths=[18, 75, 70])

    doc.add_heading("부록 D. 재현 정보", level=1)
    snapshot = summary["snapshot"]
    reproducibility_rows = [
        ["분석 스크립트", "scripts/report_single_axis_analysis.py"],
        ["문서 빌더", "scripts/build_single_axis_report.py"],
        ["분석 버전", summary["analysis_version"]],
        ["내부 부트스트랩 시드", str(summary["bootstrap_seed"])],
        ["외부 벤치마크 시드", str(summary["benchmark_seed"])],
        ["SQLite SHA-256", snapshot["snapshot_sha256"]],
        ["IOHdata ZIP SHA-256", "959c10e8397c83c56b8326d889971e3d82e9d91a552aaedc1bfe936cb241455c"],
        ["IOH package", "0.3.22"],
        ["스냅샷 최종시각", snapshot["max_kst"]],
    ]
    add_table(doc, ["항목", "값"], reproducibility_rows, widths=[45, 120])
    add_body(
        doc,
        "원시 운영 DB와 개인정보·자격증명은 vault 밖으로 복사하지 않았다. 리포트에는 집계값만 들어 있다. 공개 외부 자료의 원본 ZIP과 해시는 vault 분석 작업공간에 보존했다.",
    )


def set_headers_footers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.clear()
        footer = section.footer
        foot = footer.paragraphs[0]
        foot.clear()
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if index >= 2:
            add_page_number(foot)


def normalize_document(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Normal", "Abstract"} and paragraph.text.strip():
            paragraph.paragraph_format.widow_control = True
        for run in paragraph.runs:
            if not run.font.name:
                set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            prevent_row_split(row)


def export_source_text(doc: Document, path: Path) -> None:
    lines = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    args = parse_args()
    with (args.analysis / "analysis_summary.json").open(encoding="utf-8") as handle:
        summary = json.load(handle)
    if summary.get("external") is None:
        raise RuntimeError("analysis_summary.json does not contain external results")
    identity = preserve_identity(args.existing)
    doc = Document()
    configure_styles(doc)
    add_front_matter(doc, identity, summary)
    build_academic_report(doc, args.analysis, summary)
    add_references(doc)
    add_appendix(doc, args.analysis, summary)
    set_headers_footers(doc)
    normalize_document(doc)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    export_source_text(doc, args.source_text)
    print(args.output)


if __name__ == "__main__":
    main()
